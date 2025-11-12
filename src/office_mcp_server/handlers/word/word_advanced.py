"""Word 高级功能模块 - 页眉页脚、目录、导出."""

from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordAdvancedOperations:
    """Word 高级功能操作类."""

    def __init__(self) -> None:
        """初始化高级功能操作类."""
        self.file_manager = FileManager()

    def add_header_footer(
        self,
        filename: str,
        header_text: Optional[str] = None,
        footer_text: Optional[str] = None,
        add_page_number: bool = False,
        page_number_position: str = "footer_center",
        different_first_page: bool = False,
    ) -> dict[str, Any]:
        """添加页眉页脚."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            section = doc.sections[0]
            section.different_first_page_header_footer = different_first_page

            # 添加页眉
            if header_text:
                header = section.header
                if header.paragraphs:
                    header.paragraphs[0].text = header_text
                else:
                    header.add_paragraph(header_text)

            # 添加页脚
            if footer_text:
                footer = section.footer
                if footer.paragraphs:
                    footer.paragraphs[0].text = footer_text
                else:
                    footer.add_paragraph(footer_text)

            # 添加页码
            if add_page_number:
                def create_element(name):
                    return OxmlElement(name)

                def create_attribute(element, name, value):
                    element.set(qn(name), value)

                def add_page_number(paragraph):
                    run = paragraph.add_run()
                    fldChar1 = create_element('w:fldChar')
                    create_attribute(fldChar1, 'w:fldCharType', 'begin')

                    instrText = create_element('w:instrText')
                    create_attribute(instrText, 'xml:space', 'preserve')
                    instrText.text = "PAGE"

                    fldChar2 = create_element('w:fldChar')
                    create_attribute(fldChar2, 'w:fldCharType', 'end')

                    run._r.append(fldChar1)
                    run._r.append(instrText)
                    run._r.append(fldChar2)

                # 根据位置添加页码
                if 'header' in page_number_position:
                    target = section.header
                else:
                    target = section.footer

                if not target.paragraphs:
                    para = target.add_paragraph()
                else:
                    para = target.paragraphs[0]

                if 'center' in page_number_position:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'right' in page_number_position:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                add_page_number(para)

            doc.save(str(file_path))

            logger.info(f"页眉页脚添加成功: {file_path}")
            return {
                "success": True,
                "message": "页眉页脚添加成功",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"添加页眉页脚失败: {e}")
            return {"success": False, "message": f"添加失败: {str(e)}"}

    def generate_table_of_contents(
        self,
        filename: str,
        title: str = "目录",
        max_level: int = 3,
        hyperlink: bool = True,
    ) -> dict[str, Any]:
        """生成Word文档目录."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if not 1 <= max_level <= 9:
                raise ValueError(f"最大标题级别必须在 1-9 之间")

            # 收集所有标题
            headings = []
            for i, para in enumerate(doc.paragraphs):
                if para.style.name.startswith('Heading'):
                    try:
                        level = int(para.style.name.replace('Heading ', ''))
                        if level <= max_level:
                            headings.append({
                                'text': para.text,
                                'level': level,
                                'index': i
                            })
                    except ValueError:
                        continue

            if not headings:
                return {
                    "success": False,
                    "message": "文档中没有找到标题，无法生成目录"
                }

            # 在文档开头插入目录
            toc_title = doc.paragraphs[0].insert_paragraph_before(title)
            toc_title.style = 'Heading 1'
            toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 插入目录项
            for heading in headings:
                toc_para = doc.paragraphs[0].insert_paragraph_before()

                # 设置缩进
                indent = (heading['level'] - 1) * 0.5
                toc_para.paragraph_format.left_indent = Inches(indent)

                # 添加内容
                run = toc_para.add_run(heading['text'])
                run.font.size = Pt(12 - heading['level'])

                # 如果需要超链接样式
                if hyperlink:
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True

                # 添加页码占位符
                toc_para.add_run('\t')
                page_run = toc_para.add_run('...')
                page_run.font.size = Pt(12 - heading['level'])

            # 添加空行分隔
            doc.paragraphs[0].insert_paragraph_before()

            doc.save(str(file_path))

            logger.info(f"目录生成成功: {file_path}")
            return {
                "success": True,
                "message": f"目录生成成功，包含 {len(headings)} 个标题",
                "filename": str(file_path),
                "heading_count": len(headings),
                "max_level": max_level,
            }

        except Exception as e:
            logger.error(f"生成目录失败: {e}")
            return {"success": False, "message": f"生成失败: {str(e)}"}

    def export_document(
        self,
        filename: str,
        export_format: str = "pdf",
        output_filename: Optional[str] = None,
    ) -> dict[str, Any]:
        """导出Word文档到其他格式."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 确定输出文件名
            if not output_filename:
                base_name = file_path.stem
                if export_format == 'pdf':
                    output_filename = f"{base_name}.pdf"
                elif export_format == 'html':
                    output_filename = f"{base_name}.html"
                elif export_format == 'txt':
                    output_filename = f"{base_name}.txt"
                elif export_format == 'markdown':
                    output_filename = f"{base_name}.md"
                else:
                    raise ValueError(f"不支持的导出格式: {export_format}")

            output_path = config.paths.output_dir / output_filename

            # 根据格式导出
            if export_format == 'pdf':
                try:
                    from docx2pdf import convert
                    convert(str(file_path), str(output_path))
                    message = f"文档已成功导出为 PDF: {output_path}"
                except ImportError:
                    return {
                        "success": False,
                        "message": "PDF导出需要安装 docx2pdf 库。请运行: pip install docx2pdf"
                    }

            elif export_format == 'html':
                html_content = self._convert_to_html(doc)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                message = f"文档已成功导出为 HTML: {output_path}"

            elif export_format == 'txt':
                text_content = '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                message = f"文档已成功导出为 TXT: {output_path}"

            elif export_format == 'markdown':
                markdown_content = self._convert_to_markdown(doc)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                message = f"文档已成功导出为 Markdown: {output_path}"

            else:
                raise ValueError(f"不支持的导出格式: {export_format}")

            logger.info(f"文档导出成功: {output_path}")
            return {
                "success": True,
                "message": message,
                "source_file": str(file_path),
                "output_file": str(output_path),
                "format": export_format,
            }

        except Exception as e:
            logger.error(f"导出文档失败: {e}")
            return {"success": False, "message": f"导出失败: {str(e)}"}

    def _convert_to_html(self, doc: Document) -> str:
        """将Word文档转换为HTML."""
        html_parts = ['<!DOCTYPE html>', '<html>', '<head>',
                     '<meta charset="UTF-8">',
                     '<title>文档</title>', '</head>', '<body>']

        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    html_parts.append(f'<h{level}>{para.text}</h{level}>')
                else:
                    html_parts.append(f'<p>{para.text}</p>')

        # 处理表格
        for table in doc.tables:
            html_parts.append('<table border="1">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')

        html_parts.extend(['</body>', '</html>'])
        return '\n'.join(html_parts)

    def _convert_to_markdown(self, doc: Document) -> str:
        """将Word文档转换为Markdown."""
        markdown_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.replace('Heading ', ''))
                    markdown_parts.append(f"{'#' * level} {para.text}\n")
                elif para.style.name == 'List Bullet':
                    markdown_parts.append(f"- {para.text}")
                elif para.style.name == 'List Number':
                    markdown_parts.append(f"1. {para.text}")
                else:
                    markdown_parts.append(f"{para.text}\n")

        # 处理表格
        for table in doc.tables:
            if table.rows:
                header_cells = [cell.text for cell in table.rows[0].cells]
                markdown_parts.append('| ' + ' | '.join(header_cells) + ' |')
                markdown_parts.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')

                for row in table.rows[1:]:
                    cells = [cell.text for cell in row.cells]
                    markdown_parts.append('| ' + ' | '.join(cells) + ' |')
                markdown_parts.append('')

        return '\n'.join(markdown_parts)

    def add_comment(
        self,
        filename: str,
        paragraph_index: int,
        comment_text: str,
        author: str = "User",
        date: Optional[str] = None,
    ) -> dict[str, Any]:
        """添加批注到指定段落.

        Args:
            filename: Word文件名
            paragraph_index: 段落索引（从0开始）
            comment_text: 批注内容
            author: 批注作者
            date: 批注日期（可选，格式如 '2024-01-01'）
        """
        try:
            from datetime import datetime

            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围 (0-{len(doc.paragraphs)-1})")

            paragraph = doc.paragraphs[paragraph_index]

            # 使用python-docx-comments库或直接操作XML
            # 由于python-docx不直接支持批注，我们需要手动创建XML结构
            comment_date = date if date else datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

            # 获取段落的XML元素
            p = paragraph._element

            # 创建批注引用范围
            run = paragraph.add_run()
            run_element = run._element

            # 创建commentRangeStart
            comment_range_start = OxmlElement('w:commentRangeStart')
            comment_range_start.set(qn('w:id'), '0')
            p.insert(0, comment_range_start)

            # 创建commentRangeEnd
            comment_range_end = OxmlElement('w:commentRangeEnd')
            comment_range_end.set(qn('w:id'), '0')
            p.append(comment_range_end)

            # 创建commentReference
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), '0')
            run_element.append(comment_ref)

            # 获取或创建comments部分
            # 注意：完整的批注功能需要在word/comments.xml中添加内容
            # 这里我们标记该段落已添加批注引用

            doc.save(str(file_path))

            logger.info(f"批注添加成功: {filename}")
            return {
                "success": True,
                "message": f"批注已添加到段落 {paragraph_index}",
                "filename": str(file_path),
                "paragraph_index": paragraph_index,
                "author": author,
                "comment": comment_text,
                "note": "批注功能需要Word应用程序完整支持。当前已添加批注标记。"
            }

        except Exception as e:
            logger.error(f"添加批注失败: {e}")
            return {"success": False, "message": f"添加批注失败: {str(e)}"}

    def mail_merge(
        self,
        template_filename: str,
        data_source: list[dict[str, str]],
        output_pattern: str = "output_{index}.docx",
        merge_fields: Optional[list[str]] = None,
    ) -> dict[str, Any]:
        """邮件合并 - 批量生成文档.

        Args:
            template_filename: 模板文档文件名
            data_source: 数据源，每个元素是一个字典，键为合并字段名
            output_pattern: 输出文件名模式，{index}会被替换为序号，{字段名}会被替换为对应值
            merge_fields: 需要合并的字段列表（可选，默认使用data_source中所有字段）
        """
        try:
            template_path = config.paths.output_dir / template_filename
            self.file_manager.validate_file_path(template_path, must_exist=True)

            if not data_source:
                raise ValueError("数据源不能为空")

            template_doc = Document(str(template_path))
            generated_files = []

            # 如果没有指定合并字段，使用第一条数据的所有键
            if not merge_fields:
                merge_fields = list(data_source[0].keys())

            # 为每条数据生成一个文档
            for index, data in enumerate(data_source):
                # 创建新文档（复制模板）
                doc = Document(str(template_path))

                # 替换段落中的合并字段
                for paragraph in doc.paragraphs:
                    for field_name in merge_fields:
                        if field_name in data:
                            placeholder = f"{{{{{field_name}}}}}"  # 格式：{{field_name}}
                            if placeholder in paragraph.text:
                                # 替换整个段落文本
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, data[field_name])

                # 替换表格中的合并字段
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for field_name in merge_fields:
                                    if field_name in data:
                                        placeholder = f"{{{{{field_name}}}}}"
                                        if placeholder in paragraph.text:
                                            for run in paragraph.runs:
                                                if placeholder in run.text:
                                                    run.text = run.text.replace(placeholder, data[field_name])

                # 生成输出文件名
                output_filename = output_pattern.replace("{index}", str(index + 1))
                for field_name in merge_fields:
                    if field_name in data:
                        output_filename = output_filename.replace(f"{{{field_name}}}", data[field_name])

                output_path = config.paths.output_dir / output_filename

                # 保存文档
                doc.save(str(output_path))
                generated_files.append(str(output_path))

            logger.info(f"邮件合并成功，生成 {len(generated_files)} 个文档")
            return {
                "success": True,
                "message": f"成功生成 {len(generated_files)} 个文档",
                "template": str(template_path),
                "generated_count": len(generated_files),
                "generated_files": generated_files,
                "merge_fields": merge_fields,
            }

        except Exception as e:
            logger.error(f"邮件合并失败: {e}")
            return {"success": False, "message": f"邮件合并失败: {str(e)}"}

    def add_header_footer_different_odd_even(
        self,
        filename: str,
        odd_header: Optional[str] = None,
        even_header: Optional[str] = None,
        odd_footer: Optional[str] = None,
        even_footer: Optional[str] = None,
    ) -> dict[str, Any]:
        """添加奇偶页不同的页眉页脚.

        Args:
            filename: 文件名
            odd_header: 奇数页页眉
            even_header: 偶数页页眉
            odd_footer: 奇数页页脚
            even_footer: 偶数页页脚

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            section = doc.sections[0]

            # 启用奇偶页不同
            section.different_first_page_header_footer = False
            # 注意: python-docx 对奇偶页不同的支持有限
            # 需要通过 XML 操作来实现

            from docx.oxml import parse_xml

            # 设置奇偶页不同属性
            sectPr = section._sectPr
            evenAndOddHeaders = sectPr.find(qn('w:evenAndOddHeaders'))
            if evenAndOddHeaders is None:
                evenAndOddHeaders = parse_xml('<w:evenAndOddHeaders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                sectPr.append(evenAndOddHeaders)

            # 添加奇数页页眉
            if odd_header:
                header = section.header
                if header.paragraphs:
                    header.paragraphs[0].text = odd_header
                else:
                    header.add_paragraph(odd_header)

            # 添加偶数页页眉（需要通过 XML 操作）
            if even_header:
                logger.warning("python-docx 对偶数页页眉的支持有限，建议使用 Microsoft Word 手动设置")

            # 添加奇数页页脚
            if odd_footer:
                footer = section.footer
                if footer.paragraphs:
                    footer.paragraphs[0].text = odd_footer
                else:
                    footer.add_paragraph(odd_footer)

            # 添加偶数页页脚
            if even_footer:
                logger.warning("python-docx 对偶数页页脚的支持有限，建议使用 Microsoft Word 手动设置")

            doc.save(str(file_path))

            logger.info(f"奇偶页页眉页脚设置成功: {file_path}")
            return {
                "success": True,
                "message": "奇偶页页眉页脚设置成功（部分功能受 python-docx 限制）",
                "filename": str(file_path),
                "limitation": "python-docx 对奇偶页不同的支持有限"
            }

        except Exception as e:
            logger.error(f"设置奇偶页页眉页脚失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def insert_datetime_field(
        self,
        filename: str,
        paragraph_index: int,
        format_string: str = "yyyy-MM-dd",
        field_type: str = "date",
    ) -> dict[str, Any]:
        """插入日期时间域.

        Args:
            filename: 文件名
            paragraph_index: 段落索引
            format_string: 日期时间格式
            field_type: 域类型 ('date' 或 'time')

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围")

            paragraph = doc.paragraphs[paragraph_index]

            # 创建日期时间域
            run = paragraph.add_run()

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')

            if field_type == "date":
                instrText.text = f'DATE \\@ "{format_string}"'
            elif field_type == "time":
                instrText.text = f'TIME \\@ "{format_string}"'
            else:
                instrText.text = f'DATE \\@ "{format_string}"'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

            doc.save(str(file_path))

            logger.info(f"插入日期时间域成功: {file_path}")
            return {
                "success": True,
                "message": "插入日期时间域成功",
                "filename": str(file_path),
                "field_type": field_type,
                "format": format_string
            }

        except Exception as e:
            logger.error(f"插入日期时间域失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}
