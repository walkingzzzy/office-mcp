"""Word 基础操作模块."""

from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordBasicOperations:
    """Word 基础操作类."""

    def __init__(self) -> None:
        """初始化基础操作类."""
        self.file_manager = FileManager()

    def create_document(
        self, filename: str, title: str = "", content: str = ""
    ) -> dict[str, Any]:
        """创建 Word 文档."""
        try:
            file_path = self.file_manager.validate_file_path(filename)
            self.file_manager.validate_file_extension(filename, [".docx"])

            output_path = config.paths.output_dir / file_path.name
            self.file_manager.ensure_directory(output_path.parent)

            doc = Document()

            # 设置默认字体
            doc.styles["Normal"].font.name = config.word.default_font
            doc.styles["Normal"]._element.rPr.rFonts.set(
                qn("w:eastAsia"), config.word.default_font
            )

            # 添加标题
            if title:
                heading = doc.add_heading(title, level=1)
                heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加内容
            if content:
                paragraph = doc.add_paragraph(content)
                paragraph.paragraph_format.line_spacing = config.word.default_line_spacing

            doc.save(str(output_path))

            logger.info(f"Word 文档创建成功: {output_path}")
            return {
                "success": True,
                "message": f"Word 文档创建成功",
                "filename": str(output_path),
                "title": title,
            }

        except Exception as e:
            logger.error(f"创建 Word 文档失败: {e}")
            return {"success": False, "message": f"创建失败: {str(e)}"}

    def insert_text(
        self, filename: str, text: str, position: str = "end"
    ) -> dict[str, Any]:
        """插入文本到文档."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 插入文本
            if position == "start":
                paragraph = doc.paragraphs[0].insert_paragraph_before(text)
            else:
                paragraph = doc.add_paragraph(text)

            paragraph.paragraph_format.line_spacing = config.word.default_line_spacing

            doc.save(str(file_path))

            logger.info(f"文本插入成功: {file_path}")
            return {
                "success": True,
                "message": "文本插入成功",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"插入文本失败: {e}")
            return {"success": False, "message": f"插入失败: {str(e)}"}

    def add_heading(
        self, filename: str, text: str, level: int = 1
    ) -> dict[str, Any]:
        """添加标题."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            if not 1 <= level <= 9:
                raise ValueError(f"标题级别必须在 1-9 之间")

            doc = Document(str(file_path))
            doc.add_heading(text, level=level)
            doc.save(str(file_path))

            logger.info(f"标题添加成功: {file_path}")
            return {
                "success": True,
                "message": "标题添加成功",
                "filename": str(file_path),
                "level": level,
            }

        except Exception as e:
            logger.error(f"添加标题失败: {e}")
            return {"success": False, "message": f"添加失败: {str(e)}"}

    def add_page_break(self, filename: str) -> dict[str, Any]:
        """添加分页符."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            doc.add_page_break()
            doc.save(str(file_path))

            logger.info(f"分页符添加成功: {file_path}")
            return {
                "success": True,
                "message": "分页符添加成功",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"添加分页符失败: {e}")
            return {"success": False, "message": f"添加失败: {str(e)}"}

    def insert_image(
        self, filename: str, image_path: str, width_inches: float | None = None
    ) -> dict[str, Any]:
        """插入图片."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            img_path = Path(image_path)
            if not img_path.exists():
                raise FileNotFoundError(f"图片文件不存在: {image_path}")

            doc = Document(str(file_path))

            from docx.shared import Inches
            if width_inches:
                doc.add_picture(str(img_path), width=Inches(width_inches))
            else:
                doc.add_picture(str(img_path))

            doc.save(str(file_path))

            logger.info(f"图片插入成功: {file_path}")
            return {
                "success": True,
                "message": "图片插入成功",
                "filename": str(file_path),
                "image_path": str(img_path),
            }

        except Exception as e:
            logger.error(f"插入图片失败: {e}")
            return {"success": False, "message": f"插入失败: {str(e)}"}

    def get_document_info(self, filename: str) -> dict[str, Any]:
        """获取文档信息."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            total_text = "\n".join([p.text for p in doc.paragraphs])
            word_count = len(total_text)

            logger.info(f"获取文档信息成功: {file_path}")
            return {
                "success": True,
                "filename": str(file_path),
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "word_count": word_count,
            }

        except Exception as e:
            logger.error(f"获取文档信息失败: {e}")
            return {"success": False, "message": f"获取失败: {str(e)}"}

    def get_document_properties(self, filename: str) -> dict[str, Any]:
        """获取文档属性（元数据）.

        Returns:
            dict: 文档属性信息
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            core_props = doc.core_properties

            properties = {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "comments": core_props.comments,
                "category": core_props.category,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
            }

            logger.info(f"获取文档属性成功: {file_path}")
            return {
                "success": True,
                "message": "获取文档属性成功",
                "filename": str(file_path),
                "properties": properties
            }

        except Exception as e:
            logger.error(f"获取文档属性失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def set_document_properties(
        self,
        filename: str,
        author: Optional[str] = None,
        title: Optional[str] = None,
        subject: Optional[str] = None,
        keywords: Optional[str] = None,
        comments: Optional[str] = None,
        category: Optional[str] = None,
    ) -> dict[str, Any]:
        """设置文档属性（元数据）.

        Args:
            filename: 文件名
            author: 作者
            title: 标题
            subject: 主题
            keywords: 关键词
            comments: 备注
            category: 类别

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            core_props = doc.core_properties

            if author is not None:
                core_props.author = author
            if title is not None:
                core_props.title = title
            if subject is not None:
                core_props.subject = subject
            if keywords is not None:
                core_props.keywords = keywords
            if comments is not None:
                core_props.comments = comments
            if category is not None:
                core_props.category = category

            doc.save(str(file_path))

            logger.info(f"设置文档属性成功: {file_path}")
            return {
                "success": True,
                "message": "设置文档属性成功",
                "filename": str(file_path)
            }

        except Exception as e:
            logger.error(f"设置文档属性失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def get_page_count(self, filename: str) -> dict[str, Any]:
        """获取文档页数（估算值）.

        使用跨平台兼容的估算方法，基于文档内容（字数、段落数、表格数、图片数）估算页数。

        估算公式：
        - 基础页数 = 字数 / 每页平均字数（中文约550字/页）
        - 段落修正 = 段落数 * 0.02（每个段落约占0.02页）
        - 表格修正 = 表格数 * 0.3（每个表格约占0.3页）
        - 图片修正 = 图片数 * 0.2（每张图片约占0.2页）
        - 预估页数 = 基础页数 + 段落修正 + 表格修正 + 图片修正

        Args:
            filename: 文件名

        Returns:
            dict: 页数统计结果，包含：
                - estimated_pages: 估算的页数（整数）
                - is_estimated: true（标记为估算值）
                - estimation_basis: 估算依据
                - confidence_level: 置信度（"low"/"medium"/"high"）
                - details: 详细计算过程

        注意:
            这是估算值，实际页数可能因字体、字号、行距、段落间距、页边距等因素有所不同。
            误差范围通常在±2页以内。
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 收集文档统计信息
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # 计算字数（中文字符）
            total_text = "\n".join([p.text for p in doc.paragraphs])
            char_count = len(total_text)

            # 统计图片数量
            image_count = 0
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//w:drawing'):
                        image_count += 1

            # 估算页数
            # 中文文档：每页约550字（假设宋体12pt，1.5倍行距，标准页边距）
            chars_per_page = 550

            # 基础页数（基于字数）
            base_pages = char_count / chars_per_page if char_count > 0 else 0

            # 段落修正（每个段落约占0.02页，因为段落间距）
            paragraph_correction = paragraph_count * 0.02

            # 表格修正（每个表格约占0.3页）
            table_correction = table_count * 0.3

            # 图片修正（每张图片约占0.2页）
            image_correction = image_count * 0.2

            # 总页数估算
            total_pages = base_pages + paragraph_correction + table_correction + image_correction
            estimated_pages = max(1, round(total_pages))  # 至少1页

            # 计算置信度
            # 如果文档内容丰富（有表格、图片），置信度较低
            # 如果文档主要是纯文本，置信度较高
            if table_count > 5 or image_count > 5:
                confidence_level = "low"
            elif table_count > 0 or image_count > 0:
                confidence_level = "medium"
            else:
                confidence_level = "high"

            estimation_basis = {
                "char_count": char_count,
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "image_count": image_count,
                "chars_per_page": chars_per_page,
            }

            details = {
                "base_pages": round(base_pages, 2),
                "paragraph_correction": round(paragraph_correction, 2),
                "table_correction": round(table_correction, 2),
                "image_correction": round(image_correction, 2),
                "total_pages_raw": round(total_pages, 2),
            }

            logger.info(f"页数估算完成: {file_path}, 估算页数: {estimated_pages}")
            return {
                "success": True,
                "message": f"页数估算完成（估算值，误差±2页）",
                "filename": str(file_path),
                "estimated_pages": estimated_pages,
                "is_estimated": True,
                "confidence_level": confidence_level,
                "estimation_basis": estimation_basis,
                "details": details,
            }

        except Exception as e:
            logger.error(f"页数估算失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}
