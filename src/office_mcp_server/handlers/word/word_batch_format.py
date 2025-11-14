"""Word 批量格式化操作模块."""

from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_COLOR_INDEX
from docx.shared import Inches, Pt, RGBColor
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager
from office_mcp_server.utils.format_helper import ColorUtils


class WordBatchFormatOperations:
    """Word 批量格式化操作类."""

    def __init__(self) -> None:
        """初始化批量格式化操作类."""
        self.file_manager = FileManager()

    def batch_format_text(
        self,
        filename: str,
        paragraph_indices: list[int],
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: bool = False,
        italic: bool = False,
        color: Optional[str] = None,
        underline: Optional[str] = None,
    ) -> dict[str, Any]:
        """批量格式化文本.

        Args:
            filename: 文件名
            paragraph_indices: 段落索引列表
            font_name: 字体名称 (可选)
            font_size: 字号 (可选)
            bold: 是否加粗 (默认 False)
            italic: 是否斜体 (默认 False)
            color: 文字颜色 HEX格式 (如 '#FF0000', 可选)
            underline: 下划线样式 ('single', 'double', 'thick', 可选)

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            success_count = 0
            failed_indices = []

            for idx in paragraph_indices:
                try:
                    if idx >= len(doc.paragraphs):
                        failed_indices.append(idx)
                        continue

                    para = doc.paragraphs[idx]

                    for run in para.runs:
                        if font_name:
                            run.font.name = font_name
                        if font_size:
                            run.font.size = Pt(font_size)
                        if bold:
                            run.font.bold = True
                        if italic:
                            run.font.italic = True
                        if color:
                            r, g, b = ColorUtils.hex_to_rgb(color)
                            run.font.color.rgb = RGBColor(r, g, b)
                        if underline:
                            underline_map = {
                                'single': WD_UNDERLINE.SINGLE,
                                'double': WD_UNDERLINE.DOUBLE,
                                'thick': WD_UNDERLINE.THICK,
                            }
                            if underline in underline_map:
                                run.font.underline = underline_map[underline]

                    success_count += 1

                except Exception as e:
                    logger.warning(f"格式化段落 {idx} 失败: {e}")
                    failed_indices.append(idx)

            doc.save(str(file_path))

            logger.info(f"批量文本格式化成功: {file_path}, 成功 {success_count}/{len(paragraph_indices)} 个段落")
            return {
                "success": True,
                "message": f"成功格式化 {success_count}/{len(paragraph_indices)} 个段落",
                "filename": str(file_path),
                "success_count": success_count,
                "total_count": len(paragraph_indices),
                "failed_indices": failed_indices,
            }

        except Exception as e:
            logger.error(f"批量文本格式化失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def batch_format_paragraph(
        self,
        filename: str,
        paragraph_indices: list[int],
        alignment: Optional[str] = None,
        line_spacing: Optional[float] = None,
        space_before: Optional[float] = None,
        space_after: Optional[float] = None,
        left_indent: Optional[float] = None,
        right_indent: Optional[float] = None,
        first_line_indent: Optional[float] = None,
    ) -> dict[str, Any]:
        """批量格式化段落.

        Args:
            filename: 文件名
            paragraph_indices: 段落索引列表
            alignment: 对齐方式 ('left', 'center', 'right', 'justify', 可选)
            line_spacing: 行距倍数 (可选)
            space_before: 段前间距磅值 (可选)
            space_after: 段后间距磅值 (可选)
            left_indent: 左缩进英寸 (可选)
            right_indent: 右缩进英寸 (可选)
            first_line_indent: 首行缩进英寸 (可选)

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            success_count = 0
            failed_indices = []

            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }

            for idx in paragraph_indices:
                try:
                    if idx >= len(doc.paragraphs):
                        failed_indices.append(idx)
                        continue

                    para = doc.paragraphs[idx]
                    para_format = para.paragraph_format

                    if alignment and alignment in alignment_map:
                        para_format.alignment = alignment_map[alignment]
                    if line_spacing:
                        para_format.line_spacing = line_spacing
                    if space_before is not None:
                        para_format.space_before = Pt(space_before)
                    if space_after is not None:
                        para_format.space_after = Pt(space_after)
                    if left_indent is not None:
                        para_format.left_indent = Inches(left_indent)
                    if right_indent is not None:
                        para_format.right_indent = Inches(right_indent)
                    if first_line_indent is not None:
                        para_format.first_line_indent = Inches(first_line_indent)

                    success_count += 1

                except Exception as e:
                    logger.warning(f"格式化段落 {idx} 失败: {e}")
                    failed_indices.append(idx)

            doc.save(str(file_path))

            logger.info(f"批量段落格式化成功: {file_path}, 成功 {success_count}/{len(paragraph_indices)} 个段落")
            return {
                "success": True,
                "message": f"成功格式化 {success_count}/{len(paragraph_indices)} 个段落",
                "filename": str(file_path),
                "success_count": success_count,
                "total_count": len(paragraph_indices),
                "failed_indices": failed_indices,
            }

        except Exception as e:
            logger.error(f"批量段落格式化失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def batch_format_combined(
        self,
        filename: str,
        paragraph_indices: list[int],
        # 文本格式参数
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: bool = False,
        italic: bool = False,
        color: Optional[str] = None,
        # 段落格式参数
        alignment: Optional[str] = None,
        line_spacing: Optional[float] = None,
        space_before: Optional[float] = None,
        space_after: Optional[float] = None,
        first_line_indent: Optional[float] = None,
    ) -> dict[str, Any]:
        """批量格式化文本和段落（组合操作）.

        Args:
            filename: 文件名
            paragraph_indices: 段落索引列表
            font_name: 字体名称 (可选)
            font_size: 字号 (可选)
            bold: 是否加粗 (默认 False)
            italic: 是否斜体 (默认 False)
            color: 文字颜色 HEX格式 (可选)
            alignment: 对齐方式 (可选)
            line_spacing: 行距倍数 (可选)
            space_before: 段前间距磅值 (可选)
            space_after: 段后间距磅值 (可选)
            first_line_indent: 首行缩进英寸 (可选)

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            success_count = 0
            failed_indices = []

            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }

            for idx in paragraph_indices:
                try:
                    if idx >= len(doc.paragraphs):
                        failed_indices.append(idx)
                        continue

                    para = doc.paragraphs[idx]

                    # 格式化文本
                    for run in para.runs:
                        if font_name:
                            run.font.name = font_name
                        if font_size:
                            run.font.size = Pt(font_size)
                        if bold:
                            run.font.bold = True
                        if italic:
                            run.font.italic = True
                        if color:
                            r, g, b = ColorUtils.hex_to_rgb(color)
                            run.font.color.rgb = RGBColor(r, g, b)

                    # 格式化段落
                    para_format = para.paragraph_format
                    if alignment and alignment in alignment_map:
                        para_format.alignment = alignment_map[alignment]
                    if line_spacing:
                        para_format.line_spacing = line_spacing
                    if space_before is not None:
                        para_format.space_before = Pt(space_before)
                    if space_after is not None:
                        para_format.space_after = Pt(space_after)
                    if first_line_indent is not None:
                        para_format.first_line_indent = Inches(first_line_indent)

                    success_count += 1

                except Exception as e:
                    logger.warning(f"格式化段落 {idx} 失败: {e}")
                    failed_indices.append(idx)

            doc.save(str(file_path))

            logger.info(f"批量组合格式化成功: {file_path}, 成功 {success_count}/{len(paragraph_indices)} 个段落")
            return {
                "success": True,
                "message": f"成功格式化 {success_count}/{len(paragraph_indices)} 个段落",
                "filename": str(file_path),
                "success_count": success_count,
                "total_count": len(paragraph_indices),
                "failed_indices": failed_indices,
            }

        except Exception as e:
            logger.error(f"批量组合格式化失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

