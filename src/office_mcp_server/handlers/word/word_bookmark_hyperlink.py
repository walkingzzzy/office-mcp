"""Word 书签和超链接模块."""

from typing import Any, Optional, List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordBookmarkHyperlinkOperations:
    """Word 书签和超链接操作类."""

    def __init__(self) -> None:
        """初始化书签和超链接操作类."""
        self.file_manager = FileManager()

    def add_bookmark(
        self,
        filename: str,
        paragraph_index: int,
        bookmark_name: str,
    ) -> dict[str, Any]:
        """添加书签.

        Args:
            filename: 文件名
            paragraph_index: 段落索引
            bookmark_name: 书签名称
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围")

            paragraph = doc.paragraphs[paragraph_index]

            # 创建书签开始标记
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), '0')
            bookmark_start.set(qn('w:name'), bookmark_name)

            # 创建书签结束标记
            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), '0')

            # 插入书签标记
            paragraph._element.insert(0, bookmark_start)
            paragraph._element.append(bookmark_end)

            doc.save(str(file_path))

            logger.info(f"书签添加成功: {file_path}")
            return {
                "success": True,
                "message": f"成功添加书签 '{bookmark_name}'",
                "filename": str(file_path),
                "bookmark_name": bookmark_name,
            }

        except Exception as e:
            logger.error(f"添加书签失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def list_bookmarks(
        self,
        filename: str,
    ) -> dict[str, Any]:
        """列出所有书签.

        Args:
            filename: 文件名
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 查找所有书签
            bookmarks = []
            for element in doc.element.body.iter():
                if element.tag.endswith('bookmarkStart'):
                    bookmark_name = element.get(qn('w:name'))
                    if bookmark_name:
                        bookmarks.append(bookmark_name)

            logger.info(f"书签列表获取成功: {file_path}")
            return {
                "success": True,
                "message": f"找到 {len(bookmarks)} 个书签",
                "filename": str(file_path),
                "bookmark_count": len(bookmarks),
                "bookmarks": bookmarks,
            }

        except Exception as e:
            logger.error(f"列出书签失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def add_hyperlink(
        self,
        filename: str,
        paragraph_index: int,
        text: str,
        url: str,
        link_type: str = "url",
    ) -> dict[str, Any]:
        """添加超链接.

        Args:
            filename: 文件名
            paragraph_index: 段落索引
            text: 链接文本
            url: 链接地址
            link_type: 链接类型 ('url', 'email', 'bookmark')
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围")

            paragraph = doc.paragraphs[paragraph_index]

            # 准备链接地址
            if link_type == "email":
                full_url = f"mailto:{url}"
            elif link_type == "bookmark":
                full_url = f"#{url}"
            else:
                full_url = url

            # 添加超链接
            # 注意: python-docx 不直接支持超链接,需要通过 XML 操作
            hyperlink = self._add_hyperlink_to_paragraph(paragraph, text, full_url)

            doc.save(str(file_path))

            logger.info(f"超链接添加成功: {file_path}")
            return {
                "success": True,
                "message": "超链接添加成功",
                "filename": str(file_path),
                "text": text,
                "url": full_url,
                "link_type": link_type,
            }

        except Exception as e:
            logger.error(f"添加超链接失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def _add_hyperlink_to_paragraph(self, paragraph, text, url):
        """辅助方法: 向段落添加超链接."""
        # 获取或创建关系ID
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        # 创建超链接元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # 创建run元素
        new_run = OxmlElement('w:r')

        # 创建文本元素
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        # 设置超链接样式
        rPr = OxmlElement('w:rPr')
        # 添加蓝色下划线样式
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.insert(0, rPr)
        hyperlink.append(new_run)

        paragraph._element.append(hyperlink)

        return hyperlink

    def extract_hyperlinks(
        self,
        filename: str,
    ) -> dict[str, Any]:
        """提取所有超链接.

        Args:
            filename: 文件名
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            hyperlinks = []

            # 遍历所有段落查找超链接
            for para_idx, paragraph in enumerate(doc.paragraphs):
                for element in paragraph._element.iter():
                    if element.tag.endswith('hyperlink'):
                        # 获取链接ID
                        r_id = element.get(qn('r:id'))
                        if r_id:
                            # 获取链接目标
                            try:
                                target = paragraph.part.rels[r_id].target_ref
                                # 获取链接文本
                                text = ''.join([node.text for node in element.iter() if node.tag.endswith('t')])

                                hyperlinks.append({
                                    "paragraph_index": para_idx,
                                    "text": text,
                                    "url": target,
                                })
                            except:
                                pass

            logger.info(f"超链接提取成功: {file_path}")
            return {
                "success": True,
                "message": f"找到 {len(hyperlinks)} 个超链接",
                "filename": str(file_path),
                "hyperlink_count": len(hyperlinks),
                "hyperlinks": hyperlinks,
            }

        except Exception as e:
            logger.error(f"提取超链接失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def delete_bookmark(
        self,
        filename: str,
        bookmark_name: str,
    ) -> dict[str, Any]:
        """删除书签.

        Args:
            filename: 文件名
            bookmark_name: 书签名称

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            deleted_count = 0

            # 查找并删除书签
            for element in doc.element.body.iter():
                if element.tag.endswith('bookmarkStart'):
                    name = element.get(qn('w:name'))
                    if name == bookmark_name:
                        # 删除书签开始标记
                        element.getparent().remove(element)
                        deleted_count += 1
                elif element.tag.endswith('bookmarkEnd'):
                    # 删除对应的书签结束标记
                    # 注意: 这里简化处理，删除所有书签结束标记
                    # 实际应该匹配ID
                    pass

            if deleted_count == 0:
                return {
                    "success": False,
                    "message": f"未找到书签 '{bookmark_name}'"
                }

            doc.save(str(file_path))

            logger.info(f"删除书签成功: {file_path}, 书签: {bookmark_name}")
            return {
                "success": True,
                "message": f"成功删除书签 '{bookmark_name}'",
                "filename": str(file_path),
                "bookmark_name": bookmark_name
            }

        except Exception as e:
            logger.error(f"删除书签失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def batch_update_hyperlinks(
        self,
        filename: str,
        old_domain: str,
        new_domain: str,
    ) -> dict[str, Any]:
        """批量更新超链接中的域名.

        Args:
            filename: 文件名
            old_domain: 旧域名
            new_domain: 新域名

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            updated_count = 0

            # 遍历所有段落查找超链接
            for paragraph in doc.paragraphs:
                for element in paragraph._element.iter():
                    if element.tag.endswith('hyperlink'):
                        r_id = element.get(qn('r:id'))
                        if r_id:
                            try:
                                rel = paragraph.part.rels[r_id]
                                old_url = rel.target_ref

                                # 替换域名
                                if old_domain in old_url:
                                    new_url = old_url.replace(old_domain, new_domain)
                                    # 更新关系
                                    rel._target = new_url
                                    updated_count += 1
                            except:
                                pass

            doc.save(str(file_path))

            logger.info(f"批量更新超链接成功: {file_path}, 更新 {updated_count} 个")
            return {
                "success": True,
                "message": f"成功更新 {updated_count} 个超链接",
                "filename": str(file_path),
                "old_domain": old_domain,
                "new_domain": new_domain,
                "updated_count": updated_count
            }

        except Exception as e:
            logger.error(f"批量更新超链接失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}
