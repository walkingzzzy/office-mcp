"""Word 文档清理操作模块."""

from typing import Any, Optional

from docx import Document
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordCleanupOperations:
    """Word 文档清理操作类."""

    def __init__(self) -> None:
        """初始化清理操作类."""
        self.file_manager = FileManager()

    def delete_empty_paragraphs(
        self, filename: str
    ) -> dict[str, Any]:
        """删除所有空段落.

        Args:
            filename: 文件名

        Returns:
            dict: 操作结果，包含删除的段落数量
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            
            total_before = len(doc.paragraphs)
            deleted_count = 0
            deleted_indices = []

            # 从后向前遍历，避免索引错位
            for i in range(len(doc.paragraphs) - 1, -1, -1):
                para = doc.paragraphs[i]
                # 检查段落是否为空（去除空白字符后）
                if not para.text.strip():
                    # 删除段落元素
                    p_element = para._element
                    p_element.getparent().remove(p_element)
                    deleted_count += 1
                    deleted_indices.append(i)

            doc.save(str(file_path))

            total_after = len(doc.paragraphs)

            logger.info(f"删除空段落成功: {file_path}, 删除 {deleted_count} 个空段落")
            return {
                "success": True,
                "message": f"成功删除 {deleted_count} 个空段落",
                "filename": str(file_path),
                "deleted_count": deleted_count,
                "total_before": total_before,
                "total_after": total_after,
                "deleted_indices": sorted(deleted_indices, reverse=True),  # 降序排列
            }

        except Exception as e:
            logger.error(f"删除空段落失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def delete_paragraphs_by_indices(
        self,
        filename: str,
        paragraph_indices: list[int],
    ) -> dict[str, Any]:
        """按索引批量删除段落.

        Args:
            filename: 文件名
            paragraph_indices: 要删除的段落索引列表（从0开始）

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            
            total_before = len(doc.paragraphs)
            deleted_count = 0
            failed_indices = []

            # 对索引排序并去重，从大到小删除避免索引错位
            sorted_indices = sorted(set(paragraph_indices), reverse=True)

            for idx in sorted_indices:
                try:
                    if idx < 0 or idx >= len(doc.paragraphs):
                        failed_indices.append(idx)
                        logger.warning(f"段落索引 {idx} 超出范围，跳过")
                        continue

                    para = doc.paragraphs[idx]
                    p_element = para._element
                    p_element.getparent().remove(p_element)
                    deleted_count += 1

                except Exception as e:
                    failed_indices.append(idx)
                    logger.warning(f"删除段落 {idx} 失败: {e}")

            doc.save(str(file_path))

            total_after = len(doc.paragraphs)

            logger.info(f"批量删除段落成功: {file_path}, 删除 {deleted_count}/{len(paragraph_indices)} 个段落")
            return {
                "success": True,
                "message": f"成功删除 {deleted_count}/{len(paragraph_indices)} 个段落",
                "filename": str(file_path),
                "deleted_count": deleted_count,
                "total_requested": len(paragraph_indices),
                "total_before": total_before,
                "total_after": total_after,
                "failed_indices": failed_indices,
            }

        except Exception as e:
            logger.error(f"批量删除段落失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def analyze_page_waste(
        self, filename: str
    ) -> dict[str, Any]:
        """分析文档页面浪费情况并给出优化建议.

        Args:
            filename: 文件名

        Returns:
            dict: 分析结果和优化建议
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            analysis = {
                "empty_paragraphs": 0,
                "empty_paragraph_indices": [],
                "large_spacing": [],
                "large_font_size": [],
                "large_line_spacing": [],
                "optimization_potential_pages": 0.0,
            }

            # 遍历所有段落进行分析
            for idx, para in enumerate(doc.paragraphs):
                # 检查空段落
                if not para.text.strip():
                    analysis["empty_paragraphs"] += 1
                    analysis["empty_paragraph_indices"].append(idx)

                # 检查段落格式
                para_format = para.paragraph_format

                # 检查过大的段前段后间距（>18pt）
                if para_format.space_before and para_format.space_before.pt > 18:
                    analysis["large_spacing"].append({
                        "index": idx,
                        "text": para.text[:50] + "..." if len(para.text) > 50 else para.text,
                        "space_before": round(para_format.space_before.pt, 1),
                        "style": para.style.name,
                    })
                if para_format.space_after and para_format.space_after.pt > 18:
                    analysis["large_spacing"].append({
                        "index": idx,
                        "text": para.text[:50] + "..." if len(para.text) > 50 else para.text,
                        "space_after": round(para_format.space_after.pt, 1),
                        "style": para.style.name,
                    })

                # 检查过大的行距（>1.5倍）
                if para_format.line_spacing and para_format.line_spacing > 1.5:
                    analysis["large_line_spacing"].append({
                        "index": idx,
                        "text": para.text[:50] + "..." if len(para.text) > 50 else para.text,
                        "line_spacing": round(para_format.line_spacing, 2),
                        "style": para.style.name,
                    })

                # 检查过大的字号（>18pt）
                for run in para.runs:
                    if run.font.size and run.font.size.pt > 18:
                        analysis["large_font_size"].append({
                            "index": idx,
                            "text": para.text[:50] + "..." if len(para.text) > 50 else para.text,
                            "font_size": round(run.font.size.pt, 1),
                            "style": para.style.name,
                        })
                        break  # 每个段落只记录一次

            # 估算优化潜力（可节省的页数）
            # 每个空段落约占0.02页
            empty_para_savings = analysis["empty_paragraphs"] * 0.02
            # 每处大间距约可节省0.05页
            spacing_savings = len(analysis["large_spacing"]) * 0.05
            # 每处大字号约可节省0.03页
            font_size_savings = len(analysis["large_font_size"]) * 0.03
            # 每处大行距约可节省0.04页
            line_spacing_savings = len(analysis["large_line_spacing"]) * 0.04

            analysis["optimization_potential_pages"] = round(
                empty_para_savings + spacing_savings + font_size_savings + line_spacing_savings,
                2
            )

            # 生成优化建议
            suggestions = []
            if analysis["empty_paragraphs"] > 0:
                suggestions.append(
                    f"删除 {analysis['empty_paragraphs']} 个空段落可节省约 {empty_para_savings:.1f} 页"
                )
            if len(analysis["large_spacing"]) > 0:
                suggestions.append(
                    f"优化 {len(analysis['large_spacing'])} 处过大间距可节省约 {spacing_savings:.1f} 页"
                )
            if len(analysis["large_font_size"]) > 0:
                suggestions.append(
                    f"缩小 {len(analysis['large_font_size'])} 处过大字号可节省约 {font_size_savings:.1f} 页"
                )
            if len(analysis["large_line_spacing"]) > 0:
                suggestions.append(
                    f"减小 {len(analysis['large_line_spacing'])} 处过大行距可节省约 {line_spacing_savings:.1f} 页"
                )

            if not suggestions:
                suggestions.append("文档排版已经很紧凑，无明显优化空间")
            else:
                suggestions.append(
                    f"总计预计可节省约 {analysis['optimization_potential_pages']:.1f} 页"
                )
                suggestions.append(
                    "建议使用 auto_format_word_document 工具的 'compact' 预设进行一键优化"
                )

            logger.info(f"页面浪费分析完成: {file_path}")
            return {
                "success": True,
                "message": "页面浪费分析完成",
                "filename": str(file_path),
                "analysis": analysis,
                "suggestions": suggestions,
            }

        except Exception as e:
            logger.error(f"页面浪费分析失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def suggest_compression_strategy(
        self, filename: str
    ) -> dict[str, Any]:
        """智能推荐文档压缩策略.

        基于文档内容智能分析文档类型，并推荐最佳的压缩方案。

        Args:
            filename: 文件名

        Returns:
            dict: 压缩策略建议，包含：
                - detected_type: 检测到的文档类型
                - recommended_preset: 推荐的预设方案
                - compression_potential: 压缩潜力（"high"/"medium"/"low"）
                - specific_suggestions: 针对性的优化建议列表
                - warnings: 警告信息
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 收集文档特征
            all_text = "\n".join([p.text for p in doc.paragraphs]).lower()

            # 提取所有标题文本
            heading_texts = []
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    heading_texts.append(para.text.lower())

            heading_text = " ".join(heading_texts)

            # 文档类型检测（基于关键词）
            doc_type_keywords = {
                "商务报告": ["报告", "方案", "计划", "总结", "汇报", "提案"],
                "学术论文": ["研究", "分析", "论文", "摘要", "参考文献", "引言", "结论"],
                "技术文档": ["开发", "api", "技术", "接口", "架构", "设计", "实现", "功能"],
                "法律文书": ["合同", "协议", "条款", "甲方", "乙方", "法律", "权利", "义务"],
                "医疗报告": ["诊断", "治疗", "病历", "症状", "检查", "医疗", "患者"],
                "教育文档": ["教学", "课程", "学习", "教案", "教育", "培训", "学生"],
                "政府公文": ["通知", "公告", "决定", "意见", "办法", "规定", "文件"],
            }

            detected_type = "通用文档"
            max_score = 0

            for doc_type, keywords in doc_type_keywords.items():
                score = sum(1 for keyword in keywords if keyword in heading_text or keyword in all_text[:500])
                if score > max_score:
                    max_score = score
                    detected_type = doc_type

            # 调用页面浪费分析
            waste_analysis = self.analyze_page_waste(filename)

            if not waste_analysis["success"]:
                return waste_analysis

            optimization_potential = waste_analysis["analysis"]["optimization_potential_pages"]

            # 根据文档类型推荐预设方案
            recommendations = {
                "商务报告": {
                    "preset": "compact",
                    "reason": "商务报告适合使用紧凑排版，可有效减少页数，提升专业性",
                    "compression_potential": "high" if optimization_potential > 3 else "medium",
                },
                "学术论文": {
                    "preset": "academic",
                    "reason": "学术论文应保持标准格式，不建议过度压缩，以确保可读性和符合学术规范",
                    "compression_potential": "low",
                },
                "技术文档": {
                    "preset": "compact",
                    "reason": "技术文档适合使用紧凑排版，便于快速浏览和查阅",
                    "compression_potential": "high" if optimization_potential > 3 else "medium",
                },
                "法律文书": {
                    "preset": None,
                    "reason": "法律文书不建议压缩，应保持标准格式以确保法律效力和可读性",
                    "compression_potential": "low",
                },
                "医疗报告": {
                    "preset": "professional",
                    "reason": "医疗报告应保持专业格式，适度优化即可",
                    "compression_potential": "medium",
                },
                "教育文档": {
                    "preset": "simple",
                    "reason": "教育文档适合使用简洁风格，保持清晰易读",
                    "compression_potential": "medium",
                },
                "政府公文": {
                    "preset": None,
                    "reason": "政府公文应严格遵循公文格式规范，不建议使用预设方案",
                    "compression_potential": "low",
                },
                "通用文档": {
                    "preset": "compact",
                    "reason": "通用文档可使用紧凑排版减少页数",
                    "compression_potential": "high" if optimization_potential > 3 else "medium",
                },
            }

            recommendation = recommendations.get(detected_type, recommendations["通用文档"])

            # 生成具体建议
            specific_suggestions = []
            warnings = []

            if recommendation["preset"]:
                specific_suggestions.append(
                    f"建议使用 auto_format_word_document 工具的 '{recommendation['preset']}' 预设"
                )
            else:
                warnings.append(
                    f"⚠️ {recommendation['reason']}"
                )

            # 基于页面浪费分析添加具体建议
            if waste_analysis["analysis"]["empty_paragraphs"] > 0:
                specific_suggestions.append(
                    f"使用 delete_empty_paragraphs_in_word 删除 {waste_analysis['analysis']['empty_paragraphs']} 个空段落"
                )

            if len(waste_analysis["analysis"]["large_spacing"]) > 5:
                specific_suggestions.append(
                    "文档存在较多过大间距，建议使用紧凑排版预设优化"
                )

            if len(waste_analysis["analysis"]["large_font_size"]) > 5:
                specific_suggestions.append(
                    "文档存在较多过大字号，建议适当缩小字号"
                )

            if optimization_potential > 5:
                specific_suggestions.append(
                    f"预计可节省约 {optimization_potential:.1f} 页，压缩潜力较大"
                )
            elif optimization_potential > 2:
                specific_suggestions.append(
                    f"预计可节省约 {optimization_potential:.1f} 页，有一定压缩空间"
                )
            else:
                specific_suggestions.append(
                    "文档排版已较为紧凑，压缩空间有限"
                )

            logger.info(f"压缩策略建议完成: {file_path}, 文档类型: {detected_type}")
            return {
                "success": True,
                "message": "压缩策略建议生成完成",
                "filename": str(file_path),
                "detected_type": detected_type,
                "recommended_preset": recommendation["preset"],
                "compression_potential": recommendation["compression_potential"],
                "reason": recommendation["reason"],
                "specific_suggestions": specific_suggestions,
                "warnings": warnings,
                "optimization_potential_pages": optimization_potential,
            }

        except Exception as e:
            logger.error(f"压缩策略建议失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

