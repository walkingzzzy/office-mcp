"""Word 内容提取模块."""

from typing import Any, Optional, List

from docx import Document
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordContentExtractionOperations:
    """Word 内容提取操作类."""

    def __init__(self) -> None:
        """初始化内容提取操作类."""
        self.file_manager = FileManager()

    def extract_text(
        self,
        filename: str,
        include_tables: bool = False,
    ) -> dict[str, Any]:
        """提取所有文本.

        Args:
            filename: 文件名
            include_tables: 是否包含表格文本
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格文本
            table_texts = []
            if include_tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        table_texts.append(" | ".join(row_text))

            all_text = "\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n" + "\n".join(table_texts)

            logger.info(f"文本提取成功: {file_path}")
            return {
                "success": True,
                "message": "文本提取成功",
                "filename": str(file_path),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables) if include_tables else 0,
                "text": all_text,
                "paragraphs": paragraphs,
            }

        except Exception as e:
            logger.error(f"提取文本失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def extract_headings(
        self,
        filename: str,
        max_level: int = 9,
    ) -> dict[str, Any]:
        """提取所有标题.

        Args:
            filename: 文件名
            max_level: 最大标题级别
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            headings = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    try:
                        level = int(para.style.name.split()[-1])
                        if level <= max_level:
                            headings.append({
                                "level": level,
                                "text": para.text,
                                "style": para.style.name,
                            })
                    except (ValueError, IndexError):
                        pass

            logger.info(f"标题提取成功: {file_path}")
            return {
                "success": True,
                "message": f"提取到 {len(headings)} 个标题",
                "filename": str(file_path),
                "heading_count": len(headings),
                "headings": headings,
            }

        except Exception as e:
            logger.error(f"提取标题失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def extract_tables(
        self,
        filename: str,
    ) -> dict[str, Any]:
        """提取所有表格数据.

        Args:
            filename: 文件名
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)

                tables_data.append({
                    "table_index": table_idx,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "data": table_data,
                })

            logger.info(f"表格数据提取成功: {file_path}")
            return {
                "success": True,
                "message": f"提取到 {len(tables_data)} 个表格",
                "filename": str(file_path),
                "table_count": len(tables_data),
                "tables": tables_data,
            }

        except Exception as e:
            logger.error(f"提取表格数据失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def extract_images(
        self,
        filename: str,
    ) -> dict[str, Any]:
        """提取所有图片信息.

        Args:
            filename: 文件名
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append({
                        "image_id": rel.rId,
                        "image_type": rel.target_ref.split('.')[-1],
                        "target": rel.target_ref,
                    })

            logger.info(f"图片信息提取成功: {file_path}")
            return {
                "success": True,
                "message": f"找到 {len(images)} 张图片",
                "filename": str(file_path),
                "image_count": len(images),
                "images": images,
            }

        except Exception as e:
            logger.error(f"提取图片信息失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def get_document_statistics(
        self,
        filename: str,
    ) -> dict[str, Any]:
        """获取文档统计信息.

        Args:
            filename: 文件名
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 统计段落、表格、图片等
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # 统计字数
            total_words = 0
            total_chars = 0
            for para in doc.paragraphs:
                text = para.text
                total_words += len(text.split())
                total_chars += len(text)

            # 统计标题
            heading_count = sum(1 for para in doc.paragraphs if para.style.name.startswith('Heading'))

            # 统计图片
            image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)

            logger.info(f"文档统计信息获取成功: {file_path}")
            return {
                "success": True,
                "message": "文档统计信息获取成功",
                "filename": str(file_path),
                "statistics": {
                    "paragraph_count": paragraph_count,
                    "table_count": table_count,
                    "heading_count": heading_count,
                    "image_count": image_count,
                    "word_count": total_words,
                    "character_count": total_chars,
                }
            }

        except Exception as e:
            logger.error(f"获取文档统计信息失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}
