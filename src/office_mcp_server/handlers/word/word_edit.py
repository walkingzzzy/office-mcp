"""Word 文本编辑模块 - 查找、替换、删除等."""

from typing import Any, Optional, List
import re

from docx import Document
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordEditOperations:
    """Word 文本编辑操作类."""

    def __init__(self) -> None:
        """初始化文本编辑操作类."""
        self.file_manager = FileManager()

    def find_text(
        self,
        filename: str,
        search_text: str,
        case_sensitive: bool = False,
        whole_word: bool = False,
    ) -> dict[str, Any]:
        """查找文本.

        Args:
            filename: 文件名
            search_text: 要查找的文本
            case_sensitive: 是否区分大小写
            whole_word: 是否全字匹配
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            matches = []

            # 准备搜索模式
            pattern = re.escape(search_text)
            if whole_word:
                pattern = r'\b' + pattern + r'\b'

            flags = 0 if case_sensitive else re.IGNORECASE

            # 在段落中查找
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if re.search(pattern, paragraph.text, flags):
                    # 找到所有匹配位置
                    for match in re.finditer(pattern, paragraph.text, flags):
                        matches.append({
                            "paragraph_index": para_idx,
                            "position": match.start(),
                            "text": match.group(),
                            "context": paragraph.text
                        })

            # 在表格中查找
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            if re.search(pattern, para.text, flags):
                                matches.append({
                                    "location": "table",
                                    "table_index": table_idx,
                                    "row": row_idx,
                                    "column": cell_idx,
                                    "text": para.text
                                })

            logger.info(f"文本查找完成: {file_path}, 找到 {len(matches)} 处匹配")
            return {
                "success": True,
                "message": f"找到 {len(matches)} 处匹配",
                "filename": str(file_path),
                "search_text": search_text,
                "match_count": len(matches),
                "matches": matches
            }

        except Exception as e:
            logger.error(f"查找文本失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def replace_text(
        self,
        filename: str,
        search_text: str,
        replace_text: str,
        case_sensitive: bool = False,
        whole_word: bool = False,
        max_replacements: Optional[int] = None,
    ) -> dict[str, Any]:
        """替换文本.

        Args:
            filename: 文件名
            search_text: 要查找的文本
            replace_text: 替换为的文本
            case_sensitive: 是否区分大小写
            whole_word: 是否全字匹配
            max_replacements: 最大替换次数 (None表示全部替换)
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            replacement_count = 0

            # 准备搜索模式
            pattern = re.escape(search_text)
            if whole_word:
                pattern = r'\b' + pattern + r'\b'

            flags = 0 if case_sensitive else re.IGNORECASE

            # 在段落中替换
            for paragraph in doc.paragraphs:
                if max_replacements and replacement_count >= max_replacements:
                    break

                if re.search(pattern, paragraph.text, flags):
                    # 计算这一段会产生多少次替换
                    matches = list(re.finditer(pattern, paragraph.text, flags))
                    replacements_in_para = len(matches)

                    if max_replacements:
                        remaining = max_replacements - replacement_count
                        replacements_in_para = min(replacements_in_para, remaining)

                    # 执行替换
                    new_text = paragraph.text
                    count = 0
                    for match in matches:
                        if max_replacements and count >= replacements_in_para:
                            break
                        new_text = new_text[:match.start() + count * (len(replace_text) - len(search_text))] + \
                                   replace_text + \
                                   new_text[match.end() + count * (len(replace_text) - len(search_text)):]
                        count += 1

                    # 替换段落中的所有runs
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)

                    replacement_count += count

            # 在表格中替换
            for table in doc.tables:
                if max_replacements and replacement_count >= max_replacements:
                    break

                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if max_replacements and replacement_count >= max_replacements:
                                break

                            if re.search(pattern, para.text, flags):
                                new_text = re.sub(pattern, replace_text, para.text,
                                                count=1 if max_replacements else 0, flags=flags)

                                for run in para.runs:
                                    run.text = ''
                                if para.runs:
                                    para.runs[0].text = new_text
                                else:
                                    para.add_run(new_text)

                                replacement_count += 1

            doc.save(str(file_path))

            logger.info(f"文本替换完成: {file_path}, 替换 {replacement_count} 处")
            return {
                "success": True,
                "message": f"成功替换 {replacement_count} 处",
                "filename": str(file_path),
                "search_text": search_text,
                "replace_text": replace_text,
                "replacement_count": replacement_count
            }

        except Exception as e:
            logger.error(f"替换文本失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def delete_text(
        self,
        filename: str,
        search_text: str,
        case_sensitive: bool = False,
        whole_word: bool = False,
    ) -> dict[str, Any]:
        """删除指定文本.

        Args:
            filename: 文件名
            search_text: 要删除的文本
            case_sensitive: 是否区分大小写
            whole_word: 是否全字匹配
        """
        return self.replace_text(
            filename, search_text, "", case_sensitive, whole_word
        )

    def find_text_regex(
        self,
        filename: str,
        regex_pattern: str,
        case_sensitive: bool = False,
    ) -> dict[str, Any]:
        """使用正则表达式查找文本.

        Args:
            filename: 文件名
            regex_pattern: 正则表达式模式
            case_sensitive: 是否区分大小写

        Returns:
            dict: 查找结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            matches = []

            flags = 0 if case_sensitive else re.IGNORECASE

            # 在段落中查找
            for para_idx, paragraph in enumerate(doc.paragraphs):
                try:
                    for match in re.finditer(regex_pattern, paragraph.text, flags):
                        matches.append({
                            "paragraph_index": para_idx,
                            "position": match.start(),
                            "text": match.group(),
                            "groups": match.groups(),
                            "context": paragraph.text
                        })
                except re.error as regex_err:
                    logger.error(f"正则表达式错误: {regex_err}")
                    return {
                        "success": False,
                        "message": f"正则表达式错误: {str(regex_err)}"
                    }

            # 在表格中查找
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            try:
                                for match in re.finditer(regex_pattern, para.text, flags):
                                    matches.append({
                                        "location": "table",
                                        "table_index": table_idx,
                                        "row": row_idx,
                                        "column": cell_idx,
                                        "text": match.group(),
                                        "groups": match.groups(),
                                        "context": para.text
                                    })
                            except re.error:
                                pass

            logger.info(f"正则表达式查找完成: {file_path}, 找到 {len(matches)} 处匹配")
            return {
                "success": True,
                "message": f"找到 {len(matches)} 处匹配",
                "filename": str(file_path),
                "regex_pattern": regex_pattern,
                "match_count": len(matches),
                "matches": matches
            }

        except Exception as e:
            logger.error(f"正则表达式查找失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def replace_text_regex(
        self,
        filename: str,
        regex_pattern: str,
        replacement: str,
        case_sensitive: bool = False,
        max_replacements: Optional[int] = None,
    ) -> dict[str, Any]:
        """使用正则表达式替换文本.

        Args:
            filename: 文件名
            regex_pattern: 正则表达式模式
            replacement: 替换文本（可以使用 \\1, \\2 等引用捕获组）
            case_sensitive: 是否区分大小写
            max_replacements: 最大替换次数 (None表示全部替换)

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))
            replacement_count = 0

            flags = 0 if case_sensitive else re.IGNORECASE

            # 在段落中替换
            for paragraph in doc.paragraphs:
                if max_replacements and replacement_count >= max_replacements:
                    break

                try:
                    # 计算这一段会产生多少次替换
                    matches = list(re.finditer(regex_pattern, paragraph.text, flags))
                    if not matches:
                        continue

                    replacements_in_para = len(matches)
                    if max_replacements:
                        remaining = max_replacements - replacement_count
                        replacements_in_para = min(replacements_in_para, remaining)

                    # 执行替换
                    new_text = re.sub(
                        regex_pattern,
                        replacement,
                        paragraph.text,
                        count=replacements_in_para,
                        flags=flags
                    )

                    # 替换段落中的所有runs
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)

                    replacement_count += replacements_in_para

                except re.error as regex_err:
                    logger.error(f"正则表达式错误: {regex_err}")
                    return {
                        "success": False,
                        "message": f"正则表达式错误: {str(regex_err)}"
                    }

            # 在表格中替换
            for table in doc.tables:
                if max_replacements and replacement_count >= max_replacements:
                    break

                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if max_replacements and replacement_count >= max_replacements:
                                break

                            try:
                                matches = list(re.finditer(regex_pattern, para.text, flags))
                                if matches:
                                    new_text = re.sub(
                                        regex_pattern,
                                        replacement,
                                        para.text,
                                        count=1 if max_replacements else 0,
                                        flags=flags
                                    )

                                    for run in para.runs:
                                        run.text = ''
                                    if para.runs:
                                        para.runs[0].text = new_text
                                    else:
                                        para.add_run(new_text)

                                    replacement_count += len(matches)
                            except re.error:
                                pass

            doc.save(str(file_path))

            logger.info(f"正则表达式替换完成: {file_path}, 替换 {replacement_count} 处")
            return {
                "success": True,
                "message": f"成功替换 {replacement_count} 处",
                "filename": str(file_path),
                "regex_pattern": regex_pattern,
                "replacement": replacement,
                "replacement_count": replacement_count
            }

        except Exception as e:
            logger.error(f"正则表达式替换失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def insert_special_character(
        self,
        filename: str,
        paragraph_index: int,
        character_name: str,
        position: Optional[int] = None,
    ) -> dict[str, Any]:
        """插入特殊字符.

        Args:
            filename: 文件名
            paragraph_index: 段落索引
            character_name: 字符名称或Unicode字符
                支持的名称: 'copyright', 'registered', 'trademark',
                'degree', 'plusminus', 'micro', 'paragraph',
                'section', 'bullet', 'ellipsis', 'emdash', 'endash',
                'nbsp' (不换行空格), 'emspace', 'enspace',
                'thinspace', 'zerowidthspace'
                或直接使用Unicode字符如 '©', '®', '™' 等
            position: 插入位置（None表示末尾）

        Returns:
            dict: 操作结果
        """
        # 特殊字符映射
        special_chars = {
            'copyright': '©',
            'registered': '®',
            'trademark': '™',
            'degree': '°',
            'plusminus': '±',
            'micro': 'µ',
            'paragraph': '¶',
            'section': '§',
            'bullet': '•',
            'ellipsis': '…',
            'emdash': '—',
            'endash': '–',
            'nbsp': '\u00A0',  # 不换行空格
            'emspace': '\u2003',  # 全角空格
            'enspace': '\u2002',  # 半角空格
            'thinspace': '\u2009',  # 窄空格
            'zerowidthspace': '\u200B',  # 零宽空格
            # 数学符号
            'infinity': '∞',
            'integral': '∫',
            'sum': '∑',
            'product': '∏',
            'sqrt': '√',
            'approx': '≈',
            'notequal': '≠',
            'lessequal': '≤',
            'greaterequal': '≥',
            # 箭头
            'leftarrow': '←',
            'rightarrow': '→',
            'uparrow': '↑',
            'downarrow': '↓',
            'leftrightarrow': '↔',
            # 货币符号
            'euro': '€',
            'pound': '£',
            'yen': '¥',
            'cent': '¢',
            # 其他常用符号
            'heart': '♥',
            'spade': '♠',
            'club': '♣',
            'diamond': '♦',
            'star': '★',
            'check': '✓',
            'cross': '✗',
        }

        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index >= len(doc.paragraphs):
                return {
                    "success": False,
                    "message": f"段落索引 {paragraph_index} 超出范围"
                }

            paragraph = doc.paragraphs[paragraph_index]

            # 获取要插入的字符
            char = special_chars.get(character_name.lower(), character_name)

            # 插入字符
            if position is None:
                # 在末尾添加
                paragraph.add_run(char)
            else:
                # 在指定位置插入
                text = paragraph.text
                new_text = text[:position] + char + text[position:]

                # 清空并重新设置文本
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

            doc.save(str(file_path))

            logger.info(f"插入特殊字符完成: {file_path}, 字符: {char}")
            return {
                "success": True,
                "message": f"成功插入特殊字符: {char}",
                "filename": str(file_path),
                "character": char,
                "paragraph_index": paragraph_index
            }

        except Exception as e:
            logger.error(f"插入特殊字符失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}
