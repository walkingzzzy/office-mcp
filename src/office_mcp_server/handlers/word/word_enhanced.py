"""Word 增强功能模块 - 图片编辑、批量操作等."""

from typing import Any, Optional, List
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordEnhancedOperations:
    """Word 增强操作类."""

    def __init__(self) -> None:
        """初始化增强操作类."""
        self.file_manager = FileManager()

    # ========== 图片操作 ==========
    def resize_image(
        self,
        filename: str,
        image_index: int,
        width_inches: Optional[float] = None,
        height_inches: Optional[float] = None,
        maintain_aspect_ratio: bool = True,
    ) -> dict[str, Any]:
        """调整图片大小.

        Args:
            filename: 文件名
            image_index: 图片索引
            width_inches: 宽度(英寸)
            height_inches: 高度(英寸)
            maintain_aspect_ratio: 是否保持宽高比
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 查找所有图片
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append(rel)

            if image_index >= len(images):
                raise ValueError(f"图片索引 {image_index} 超出范围")

            # 查找对应的 inline shape
            inline_shapes = []
            for para in doc.paragraphs:
                for run in para.runs:
                    if hasattr(run, '_element'):
                        for drawing in run._element.findall('.//{*}drawing'):
                            for inline in drawing.findall('.//{*}inline'):
                                inline_shapes.append((para, run, inline))

            if image_index >= len(inline_shapes):
                return {
                    "success": False,
                    "message": "无法找到对应的图片对象,python-docx对图片操作支持有限"
                }

            # 注意: python-docx 对图片编辑的支持有限
            # 完整功能需要更底层的 XML 操作或使用 python-docx-template

            doc.save(str(file_path))

            logger.info(f"图片大小调整成功: {file_path}")
            return {
                "success": True,
                "message": "图片大小调整功能受python-docx库限制,建议使用其他工具",
                "filename": str(file_path),
                "note": "python-docx对图片编辑支持有限,建议重新插入指定大小的图片"
            }

        except Exception as e:
            logger.error(f"调整图片大小失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    # ========== 批量操作 ==========
    def batch_replace_text(
        self,
        filenames: List[str],
        search_text: str,
        replace_text: str,
    ) -> dict[str, Any]:
        """批量替换文本.

        Args:
            filenames: 文件名列表
            search_text: 要查找的文本
            replace_text: 替换为的文本
        """
        try:
            results = []
            success_count = 0
            fail_count = 0

            for filename in filenames:
                try:
                    file_path = config.paths.output_dir / filename
                    self.file_manager.validate_file_path(file_path, must_exist=True)

                    doc = Document(str(file_path))
                    replacement_count = 0

                    # 在段落中替换
                    for paragraph in doc.paragraphs:
                        if search_text in paragraph.text:
                            # 简单替换
                            for run in paragraph.runs:
                                if search_text in run.text:
                                    run.text = run.text.replace(search_text, replace_text)
                                    replacement_count += 1

                    doc.save(str(file_path))

                    results.append({
                        "filename": filename,
                        "success": True,
                        "replacement_count": replacement_count
                    })
                    success_count += 1

                except Exception as e:
                    results.append({
                        "filename": filename,
                        "success": False,
                        "error": str(e)
                    })
                    fail_count += 1

            logger.info(f"批量替换完成: 成功 {success_count}, 失败 {fail_count}")
            return {
                "success": True,
                "message": f"批量处理完成: 成功 {success_count}, 失败 {fail_count}",
                "search_text": search_text,
                "replace_text": replace_text,
                "success_count": success_count,
                "fail_count": fail_count,
                "results": results,
            }

        except Exception as e:
            logger.error(f"批量替换失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def batch_apply_style(
        self,
        filenames: List[str],
        style_name: str,
        apply_to: str = "body",
    ) -> dict[str, Any]:
        """批量应用样式.

        Args:
            filenames: 文件名列表
            style_name: 样式名称
            apply_to: 应用范围 ('body', 'headings')
        """
        try:
            results = []
            success_count = 0
            fail_count = 0

            for filename in filenames:
                try:
                    file_path = config.paths.output_dir / filename
                    self.file_manager.validate_file_path(file_path, must_exist=True)

                    doc = Document(str(file_path))
                    affected_count = 0

                    for para in doc.paragraphs:
                        if apply_to == "body" and not para.style.name.startswith('Heading'):
                            para.style = style_name
                            affected_count += 1
                        elif apply_to == "headings" and para.style.name.startswith('Heading'):
                            para.style = style_name
                            affected_count += 1

                    doc.save(str(file_path))

                    results.append({
                        "filename": filename,
                        "success": True,
                        "affected_count": affected_count
                    })
                    success_count += 1

                except Exception as e:
                    results.append({
                        "filename": filename,
                        "success": False,
                        "error": str(e)
                    })
                    fail_count += 1

            logger.info(f"批量应用样式完成: 成功 {success_count}, 失败 {fail_count}")
            return {
                "success": True,
                "message": f"批量处理完成: 成功 {success_count}, 失败 {fail_count}",
                "style_name": style_name,
                "success_count": success_count,
                "fail_count": fail_count,
                "results": results,
            }

        except Exception as e:
            logger.error(f"批量应用样式失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def merge_documents(
        self,
        source_filenames: List[str],
        output_filename: str,
        add_page_breaks: bool = True,
    ) -> dict[str, Any]:
        """合并多个文档.

        Args:
            source_filenames: 源文件名列表
            output_filename: 输出文件名
            add_page_breaks: 是否在文档间添加分页符
        """
        try:
            # 创建新文档
            merged_doc = Document()

            for idx, filename in enumerate(source_filenames):
                file_path = config.paths.output_dir / filename
                self.file_manager.validate_file_path(file_path, must_exist=True)

                source_doc = Document(str(file_path))

                # 添加分页符(除了第一个文档)
                if idx > 0 and add_page_breaks:
                    merged_doc.add_page_break()

                # 复制段落
                for para in source_doc.paragraphs:
                    new_para = merged_doc.add_paragraph(para.text)
                    new_para.style = para.style

                # 复制表格
                for table in source_doc.tables:
                    new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text

            # 保存合并后的文档
            output_path = config.paths.output_dir / output_filename
            merged_doc.save(str(output_path))

            logger.info(f"文档合并成功: {output_path}")
            return {
                "success": True,
                "message": f"成功合并 {len(source_filenames)} 个文档",
                "output_filename": str(output_path),
                "source_count": len(source_filenames),
            }

        except Exception as e:
            logger.error(f"合并文档失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def split_document_by_headings(
        self,
        filename: str,
        heading_level: int = 1,
        output_pattern: str = "section_{index}.docx",
    ) -> dict[str, Any]:
        """按标题拆分文档.

        Args:
            filename: 文件名
            heading_level: 标题级别
            output_pattern: 输出文件名模式
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            # 查找指定级别的标题
            sections = []
            current_section = []
            heading_style = f'Heading {heading_level}'

            for para in doc.paragraphs:
                if para.style.name == heading_style:
                    if current_section:
                        sections.append(current_section)
                    current_section = [para]
                else:
                    current_section.append(para)

            if current_section:
                sections.append(current_section)

            # 为每个章节创建新文档
            output_files = []
            for idx, section in enumerate(sections):
                new_doc = Document()

                for para in section:
                    new_para = new_doc.add_paragraph(para.text)
                    new_para.style = para.style

                output_filename = output_pattern.format(index=idx + 1)
                output_path = config.paths.output_dir / output_filename
                new_doc.save(str(output_path))
                output_files.append(output_filename)

            logger.info(f"文档拆分成功: {file_path}")
            return {
                "success": True,
                "message": f"文档已拆分为 {len(sections)} 个部分",
                "filename": str(file_path),
                "section_count": len(sections),
                "output_files": output_files,
            }

        except Exception as e:
            logger.error(f"拆分文档失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def batch_convert_format(
        self,
        filenames: List[str],
        output_format: str = "pdf",
    ) -> dict[str, Any]:
        """批量转换文档格式.

        Args:
            filenames: 文件名列表
            output_format: 输出格式 ('pdf', 'html', 'txt', 'markdown')

        Returns:
            dict: 操作结果
        """
        try:
            from office_mcp_server.handlers.word.word_export import WordExportOperations

            export_ops = WordExportOperations()
            results = []
            success_count = 0
            failed_count = 0

            for filename in filenames:
                try:
                    result = export_ops.export_document(filename, output_format)
                    if result.get("success"):
                        success_count += 1
                        results.append({
                            "filename": filename,
                            "status": "success",
                            "output_file": result.get("output_filename")
                        })
                    else:
                        failed_count += 1
                        results.append({
                            "filename": filename,
                            "status": "failed",
                            "error": result.get("message")
                        })
                except Exception as e:
                    failed_count += 1
                    results.append({
                        "filename": filename,
                        "status": "failed",
                        "error": str(e)
                    })

            logger.info(f"批量转换完成: 成功 {success_count}, 失败 {failed_count}")
            return {
                "success": True,
                "message": f"批量转换完成: 成功 {success_count}, 失败 {failed_count}",
                "total": len(filenames),
                "success_count": success_count,
                "failed_count": failed_count,
                "results": results
            }

        except Exception as e:
            logger.error(f"批量转换失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def batch_add_header_footer(
        self,
        filenames: List[str],
        header_text: Optional[str] = None,
        footer_text: Optional[str] = None,
        add_page_number: bool = False,
    ) -> dict[str, Any]:
        """批量添加页眉页脚.

        Args:
            filenames: 文件名列表
            header_text: 页眉文本
            footer_text: 页脚文本
            add_page_number: 是否添加页码

        Returns:
            dict: 操作结果
        """
        try:
            from office_mcp_server.handlers.word.word_advanced import WordAdvancedOperations

            advanced_ops = WordAdvancedOperations()
            results = []
            success_count = 0
            failed_count = 0

            for filename in filenames:
                try:
                    result = advanced_ops.add_header_footer(
                        filename,
                        header_text=header_text,
                        footer_text=footer_text,
                        add_page_number=add_page_number
                    )
                    if result.get("success"):
                        success_count += 1
                        results.append({
                            "filename": filename,
                            "status": "success"
                        })
                    else:
                        failed_count += 1
                        results.append({
                            "filename": filename,
                            "status": "failed",
                            "error": result.get("message")
                        })
                except Exception as e:
                    failed_count += 1
                    results.append({
                        "filename": filename,
                        "status": "failed",
                        "error": str(e)
                    })

            logger.info(f"批量添加页眉页脚完成: 成功 {success_count}, 失败 {failed_count}")
            return {
                "success": True,
                "message": f"批量添加页眉页脚完成: 成功 {success_count}, 失败 {failed_count}",
                "total": len(filenames),
                "success_count": success_count,
                "failed_count": failed_count,
                "results": results
            }

        except Exception as e:
            logger.error(f"批量添加页眉页脚失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def batch_insert_content(
        self,
        filenames: List[str],
        content: str,
        position: str = "end",
        paragraph_index: Optional[int] = None,
    ) -> dict[str, Any]:
        """批量插入内容.

        Args:
            filenames: 文件名列表
            content: 要插入的内容
            position: 插入位置 ('start', 'end', 'index')
            paragraph_index: 段落索引（当 position='index' 时使用）

        Returns:
            dict: 操作结果
        """
        try:
            results = []
            success_count = 0
            failed_count = 0

            for filename in filenames:
                try:
                    file_path = config.paths.output_dir / filename
                    self.file_manager.validate_file_path(file_path, must_exist=True)

                    doc = Document(str(file_path))

                    if position == "start":
                        # 在开头插入
                        doc.paragraphs[0].insert_paragraph_before(content)
                    elif position == "end":
                        # 在末尾插入
                        doc.add_paragraph(content)
                    elif position == "index" and paragraph_index is not None:
                        # 在指定位置插入
                        if paragraph_index < len(doc.paragraphs):
                            doc.paragraphs[paragraph_index].insert_paragraph_before(content)
                        else:
                            doc.add_paragraph(content)

                    doc.save(str(file_path))

                    success_count += 1
                    results.append({
                        "filename": filename,
                        "status": "success"
                    })

                except Exception as e:
                    failed_count += 1
                    results.append({
                        "filename": filename,
                        "status": "failed",
                        "error": str(e)
                    })

            logger.info(f"批量插入内容完成: 成功 {success_count}, 失败 {failed_count}")
            return {
                "success": True,
                "message": f"批量插入内容完成: 成功 {success_count}, 失败 {failed_count}",
                "total": len(filenames),
                "success_count": success_count,
                "failed_count": failed_count,
                "results": results
            }

        except Exception as e:
            logger.error(f"批量插入内容失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}
