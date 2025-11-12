"""Word 格式化操作模块."""

from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager
from office_mcp_server.utils.format_helper import ColorUtils


class WordFormatOperations:
    """Word 格式化操作类."""

    def __init__(self) -> None:
        """初始化格式化操作类."""
        self.file_manager = FileManager()

    def format_text(
        self,
        filename: str,
        paragraph_index: int,
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: bool = False,
        italic: bool = False,
        color: Optional[str] = None,
        underline: Optional[str] = None,
        strike: bool = False,
        double_strike: bool = False,
        superscript: bool = False,
        subscript: bool = False,
        highlight: Optional[str] = None,
        spacing: Optional[float] = None,
        shadow: bool = False,
    ) -> dict[str, Any]:
        """格式化文本."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围")

            paragraph = doc.paragraphs[paragraph_index]

            # 下划线样式映射
            underline_styles = {
                'single': WD_UNDERLINE.SINGLE,
                'double': WD_UNDERLINE.DOUBLE,
                'thick': WD_UNDERLINE.THICK,
                'dotted': WD_UNDERLINE.DOTTED,
                'dash': WD_UNDERLINE.DASH,
                'wave': WD_UNDERLINE.WAVY,
            }

            # 高亮颜色映射
            highlight_colors = {
                'yellow': WD_COLOR_INDEX.YELLOW,
                'green': WD_COLOR_INDEX.BRIGHT_GREEN,
                'cyan': WD_COLOR_INDEX.TURQUOISE,
                'magenta': WD_COLOR_INDEX.PINK,
                'blue': WD_COLOR_INDEX.BLUE,
                'red': WD_COLOR_INDEX.RED,
                'darkBlue': WD_COLOR_INDEX.DARK_BLUE,
                'darkCyan': WD_COLOR_INDEX.TEAL,
                'darkGreen': WD_COLOR_INDEX.GREEN,
                'darkMagenta': WD_COLOR_INDEX.VIOLET,
                'darkRed': WD_COLOR_INDEX.DARK_RED,
                'darkYellow': WD_COLOR_INDEX.DARK_YELLOW,
                'darkGray': WD_COLOR_INDEX.GRAY_25,
                'lightGray': WD_COLOR_INDEX.GRAY_50,
                'black': WD_COLOR_INDEX.BLACK,
            }

            # 应用格式
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

                if font_size:
                    run.font.size = Pt(font_size)

                run.font.bold = bold
                run.font.italic = italic

                if color:
                    r, g, b = ColorUtils.hex_to_rgb(color)
                    run.font.color.rgb = RGBColor(r, g, b)

                if underline and underline in underline_styles:
                    run.font.underline = underline_styles[underline]

                run.font.strike = strike
                run.font.double_strike = double_strike
                run.font.superscript = superscript
                run.font.subscript = subscript

                if highlight and highlight in highlight_colors:
                    run.font.highlight_color = highlight_colors[highlight]

                if spacing is not None:
                    run.font.spacing = Pt(spacing)

                run.font.shadow = shadow

            doc.save(str(file_path))

            logger.info(f"文本格式化成功: {file_path}")
            return {
                "success": True,
                "message": "文本格式化成功",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"格式化文本失败: {e}")
            return {"success": False, "message": f"格式化失败: {str(e)}"}

    def format_paragraph(
        self,
        filename: str,
        paragraph_index: int,
        alignment: Optional[str] = None,
        line_spacing: Optional[float] = None,
        space_before: Optional[float] = None,
        space_after: Optional[float] = None,
        left_indent: Optional[float] = None,
        right_indent: Optional[float] = None,
        first_line_indent: Optional[float] = None,
    ) -> dict[str, Any]:
        """格式化段落."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围")

            paragraph = doc.paragraphs[paragraph_index]
            fmt = paragraph.paragraph_format

            # 对齐方式
            if alignment:
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                if alignment in alignment_map:
                    fmt.alignment = alignment_map[alignment]

            # 行距
            if line_spacing is not None:
                fmt.line_spacing = line_spacing

            # 段落间距
            if space_before is not None:
                fmt.space_before = Pt(space_before)
            if space_after is not None:
                fmt.space_after = Pt(space_after)

            # 缩进
            if left_indent is not None:
                fmt.left_indent = Inches(left_indent)
            if right_indent is not None:
                fmt.right_indent = Inches(right_indent)
            if first_line_indent is not None:
                fmt.first_line_indent = Inches(first_line_indent)

            doc.save(str(file_path))

            logger.info(f"段落格式化成功: {file_path}")
            return {
                "success": True,
                "message": "段落格式化成功",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"格式化段落失败: {e}")
            return {"success": False, "message": f"格式化失败: {str(e)}"}

    def apply_style(
        self, filename: str, paragraph_index: int, style_name: str
    ) -> dict[str, Any]:
        """应用样式到段落."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围")

            paragraph = doc.paragraphs[paragraph_index]
            paragraph.style = style_name

            doc.save(str(file_path))

            logger.info(f"样式应用成功: {file_path}")
            return {
                "success": True,
                "message": "样式应用成功",
                "filename": str(file_path),
                "style": style_name,
            }

        except Exception as e:
            logger.error(f"应用样式失败: {e}")
            return {"success": False, "message": f"应用失败: {str(e)}"}
