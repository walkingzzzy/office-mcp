"""Word文档格式检查器."""

from typing import Any, Optional
from docx import Document
from docx.shared import RGBColor
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordFormatInspector:
    """Word文档格式检查器类."""

    def __init__(self):
        """初始化格式检查器."""
        self.file_manager = FileManager()

    def get_paragraph_format(
        self,
        filename: str,
        paragraph_index: int,
    ) -> dict[str, Any]:
        """获取指定段落的详细格式信息.

        Args:
            filename: 文件名
            paragraph_index: 段落索引（从0开始）

        Returns:
            dict: 段落格式信息
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                raise ValueError(f"段落索引 {paragraph_index} 超出范围 (0-{len(doc.paragraphs)-1})")

            para = doc.paragraphs[paragraph_index]

            # 获取字体信息（从第一个run获取，如果有的话）
            font_info = {}
            if para.runs:
                first_run = para.runs[0]
                font_info = {
                    "name": first_run.font.name,
                    "size": first_run.font.size.pt if first_run.font.size else None,
                    "bold": first_run.font.bold,
                    "italic": first_run.font.italic,
                    "underline": first_run.font.underline,
                    "color": self._rgb_to_hex(first_run.font.color.rgb) if first_run.font.color.rgb else None,
                }

            # 获取段落格式信息
            pf = para.paragraph_format
            paragraph_format = {
                "alignment": str(pf.alignment) if pf.alignment else None,
                "line_spacing": pf.line_spacing,
                "space_before": pf.space_before.pt if pf.space_before else None,
                "space_after": pf.space_after.pt if pf.space_after else None,
                "left_indent": pf.left_indent.inches if pf.left_indent else None,
                "right_indent": pf.right_indent.inches if pf.right_indent else None,
                "first_line_indent": pf.first_line_indent.inches if pf.first_line_indent else None,
            }

            logger.info(f"获取段落格式成功: {filename}, 段落 {paragraph_index}")
            return {
                "success": True,
                "paragraph_index": paragraph_index,
                "text": para.text[:100] + "..." if len(para.text) > 100 else para.text,
                "style": para.style.name,
                "font": font_info,
                "paragraph_format": paragraph_format,
            }

        except Exception as e:
            logger.error(f"获取段落格式失败: {e}")
            return {"success": False, "message": f"获取失败: {str(e)}"}

    def check_document_formatting(
        self,
        filename: str,
        check_items: Optional[list[str]] = None,
    ) -> dict[str, Any]:
        """批量检查文档所有段落的格式一致性.

        Args:
            filename: 文件名
            check_items: 要检查的项目列表（可选），如["font", "alignment", "spacing"]

        Returns:
            dict: 格式检查结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if check_items is None:
                check_items = ["font", "alignment", "spacing"]

            # 收集所有段落的格式信息
            fonts_used = set()
            font_sizes_used = set()
            alignments_used = set()
            inconsistencies = []

            for i, para in enumerate(doc.paragraphs):
                if para.runs:
                    for run in para.runs:
                        if run.font.name:
                            fonts_used.add(run.font.name)
                        if run.font.size:
                            font_sizes_used.add(run.font.size.pt)

                if para.paragraph_format.alignment:
                    alignments_used.add(str(para.paragraph_format.alignment))

            # 检查不一致性
            if "font" in check_items and len(fonts_used) > 3:
                inconsistencies.append({
                    "type": "font_variety",
                    "description": f"文档使用了{len(fonts_used)}种不同的字体，可能影响一致性",
                    "fonts": list(fonts_used),
                })

            if "font" in check_items and len(font_sizes_used) > 5:
                inconsistencies.append({
                    "type": "font_size_variety",
                    "description": f"文档使用了{len(font_sizes_used)}种不同的字号",
                    "sizes": sorted(list(font_sizes_used)),
                })

            logger.info(f"文档格式检查完成: {filename}")
            return {
                "success": True,
                "total_paragraphs": len(doc.paragraphs),
                "format_summary": {
                    "fonts_used": list(fonts_used),
                    "font_sizes_used": sorted(list(font_sizes_used)),
                    "alignments_used": list(alignments_used),
                    "inconsistencies": inconsistencies,
                },
            }

        except Exception as e:
            logger.error(f"文档格式检查失败: {e}")
            return {"success": False, "message": f"检查失败: {str(e)}"}

    def get_table_format(
        self,
        filename: str,
        table_index: int,
    ) -> dict[str, Any]:
        """获取表格的格式信息.

        Args:
            filename: 文件名
            table_index: 表格索引（从0开始）

        Returns:
            dict: 表格格式信息
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index < 0 or table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围 (0-{len(doc.tables)-1})")

            table = doc.tables[table_index]

            # 获取表格样式
            style = table.style.name if table.style else None

            # 获取部分单元格格式（前3行）
            cell_formats = []
            for i, row in enumerate(table.rows[:3]):
                for j, cell in enumerate(row.cells):
                    cell_info = {
                        "row": i,
                        "col": j,
                        "text": cell.text[:50] + "..." if len(cell.text) > 50 else cell.text,
                    }

                    # 获取单元格内第一个段落的格式
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        if para.runs:
                            run = para.runs[0]
                            cell_info["font"] = {
                                "name": run.font.name,
                                "size": run.font.size.pt if run.font.size else None,
                                "bold": run.font.bold,
                            }
                        cell_info["alignment"] = str(para.paragraph_format.alignment) if para.paragraph_format.alignment else None

                    cell_formats.append(cell_info)

            logger.info(f"获取表格格式成功: {filename}, 表格 {table_index}")
            return {
                "success": True,
                "table_index": table_index,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "style": style,
                "cell_formats": cell_formats,
            }

        except Exception as e:
            logger.error(f"获取表格格式失败: {e}")
            return {"success": False, "message": f"获取失败: {str(e)}"}

    def _rgb_to_hex(self, rgb: RGBColor) -> str:
        """将RGB颜色转换为HEX格式.

        Args:
            rgb: RGBColor对象

        Returns:
            str: HEX颜色字符串（如 '#FF0000'）
        """
        if rgb is None:
            return None
        return f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

