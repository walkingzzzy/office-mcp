"""Word 图片操作模块 - 图片插入和格式化."""

from typing import Any, Optional
from pathlib import Path
import io
import requests

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordImageOperations:
    """Word 图片操作类."""

    def __init__(self) -> None:
        """初始化图片操作类."""
        self.file_manager = FileManager()

    def insert_image_from_url(
        self,
        filename: str,
        image_url: str,
        width_inches: Optional[float] = None,
        height_inches: Optional[float] = None,
        alignment: str = "left",
    ) -> dict[str, Any]:
        """从 URL 插入图片.

        Args:
            filename: 文件名
            image_url: 图片 URL
            width_inches: 图片宽度（英寸）
            height_inches: 图片高度（英寸）
            alignment: 对齐方式 ('left', 'center', 'right')

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            # 下载图片
            logger.info(f"正在从 URL 下载图片: {image_url}")
            response = requests.get(image_url, timeout=30)
            response.raise_for_status()

            # 创建图片流
            image_stream = io.BytesIO(response.content)

            doc = Document(str(file_path))

            # 添加段落用于放置图片
            paragraph = doc.add_paragraph()

            # 设置对齐方式
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
            }
            paragraph.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

            # 插入图片
            run = paragraph.add_run()
            if width_inches and height_inches:
                run.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))
            elif width_inches:
                run.add_picture(image_stream, width=Inches(width_inches))
            elif height_inches:
                run.add_picture(image_stream, height=Inches(height_inches))
            else:
                run.add_picture(image_stream)

            doc.save(str(file_path))

            logger.info(f"从 URL 插入图片成功: {file_path}")
            return {
                "success": True,
                "message": "从 URL 插入图片成功",
                "filename": str(file_path),
                "image_url": image_url,
                "alignment": alignment
            }

        except requests.RequestException as e:
            logger.error(f"下载图片失败: {e}")
            return {"success": False, "message": f"下载图片失败: {str(e)}"}
        except Exception as e:
            logger.error(f"插入图片失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def insert_image_with_size(
        self,
        filename: str,
        image_path: str,
        width_inches: Optional[float] = None,
        height_inches: Optional[float] = None,
        alignment: str = "left",
        keep_aspect_ratio: bool = True,
    ) -> dict[str, Any]:
        """插入图片并设置完整的大小和对齐方式.

        Args:
            filename: 文件名
            image_path: 图片路径
            width_inches: 图片宽度（英寸）
            height_inches: 图片高度（英寸）
            alignment: 对齐方式 ('left', 'center', 'right')
            keep_aspect_ratio: 是否保持宽高比（如果只指定宽度或高度）

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            img_path = Path(image_path)
            if not img_path.exists():
                raise FileNotFoundError(f"图片文件不存在: {image_path}")

            doc = Document(str(file_path))

            # 添加段落用于放置图片
            paragraph = doc.add_paragraph()

            # 设置对齐方式
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
            }
            paragraph.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

            # 插入图片
            run = paragraph.add_run()
            if width_inches and height_inches:
                picture = run.add_picture(str(img_path), width=Inches(width_inches), height=Inches(height_inches))
            elif width_inches:
                picture = run.add_picture(str(img_path), width=Inches(width_inches))
            elif height_inches:
                picture = run.add_picture(str(img_path), height=Inches(height_inches))
            else:
                picture = run.add_picture(str(img_path))

            doc.save(str(file_path))

            logger.info(f"插入图片成功: {file_path}")
            return {
                "success": True,
                "message": "插入图片成功",
                "filename": str(file_path),
                "image_path": str(img_path),
                "width_inches": width_inches,
                "height_inches": height_inches,
                "alignment": alignment
            }

        except Exception as e:
            logger.error(f"插入图片失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

