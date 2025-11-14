"""Word 结构操作模块 - 表格和列表."""

from typing import Any, Optional

from docx import Document
from docx.shared import Inches
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager


class WordStructureOperations:
    """Word 结构操作类."""

    def __init__(self) -> None:
        """初始化结构操作类."""
        self.file_manager = FileManager()

    def create_table(
        self, filename: str, rows: int, cols: int, data: Optional[list[list[str]]] = None
    ) -> dict[str, Any]:
        """创建表格."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"

            # 填充数据
            if data:
                for i, row_data in enumerate(data):
                    if i >= rows:
                        break
                    for j, cell_data in enumerate(row_data):
                        if j >= cols:
                            break
                        table.rows[i].cells[j].text = str(cell_data)

            doc.save(str(file_path))

            logger.info(f"表格创建成功: {file_path}")
            return {
                "success": True,
                "message": "表格创建成功",
                "filename": str(file_path),
                "rows": rows,
                "cols": cols,
            }

        except Exception as e:
            logger.error(f"创建表格失败: {e}")
            return {"success": False, "message": f"创建失败: {str(e)}"}

    def edit_table(
        self,
        filename: str,
        table_index: int,
        operation: str,
        row_index: Optional[int] = None,
        col_index: Optional[int] = None,
    ) -> dict[str, Any]:
        """编辑表格(插入/删除行列)."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            # 执行操作
            if operation == 'add_row':
                row = table.add_row()
                message = f"成功在表格末尾添加新行"

            elif operation == 'delete_row':
                if row_index is None or row_index >= len(table.rows):
                    raise ValueError(f"无效的行索引")
                table._element.remove(table.rows[row_index]._element)
                message = f"成功删除第 {row_index} 行"

            elif operation == 'add_column':
                for row in table.rows:
                    row.add_cell()
                message = f"成功在表格末尾添加新列"

            elif operation == 'delete_column':
                if col_index is None:
                    raise ValueError(f"删除列需要指定列索引")
                for row in table.rows:
                    if col_index < len(row.cells):
                        cell = row.cells[col_index]
                        cell._element.getparent().remove(cell._element)
                message = f"成功删除第 {col_index} 列"

            else:
                raise ValueError(f"不支持的操作类型: {operation}")

            doc.save(str(file_path))

            logger.info(f"表格编辑成功: {file_path}")
            return {
                "success": True,
                "message": message,
                "filename": str(file_path),
                "operation": operation,
            }

        except Exception as e:
            logger.error(f"编辑表格失败: {e}")
            return {"success": False, "message": f"编辑失败: {str(e)}"}

    def merge_table_cells(
        self,
        filename: str,
        table_index: int,
        start_row: int,
        start_col: int,
        end_row: int,
        end_col: int,
    ) -> dict[str, Any]:
        """合并表格单元格."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            if (start_row >= len(table.rows) or end_row >= len(table.rows) or
                start_col >= len(table.columns) or end_col >= len(table.columns)):
                raise ValueError(f"单元格范围超出表格边界")

            if start_row > end_row or start_col > end_col:
                raise ValueError(f"起始位置必须在结束位置之前")

            start_cell = table.cell(start_row, start_col)
            end_cell = table.cell(end_row, end_col)
            start_cell.merge(end_cell)

            doc.save(str(file_path))

            logger.info(f"单元格合并成功: {file_path}")
            return {
                "success": True,
                "message": "单元格合并成功",
                "filename": str(file_path),
                "merged_range": f"({start_row},{start_col}) 到 ({end_row},{end_col})",
            }

        except Exception as e:
            logger.error(f"合并单元格失败: {e}")
            return {"success": False, "message": f"合并失败: {str(e)}"}

    def add_list_paragraph(
        self, filename: str, text: str, list_type: str = "bullet", level: int = 0
    ) -> dict[str, Any]:
        """添加列表段落."""
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            if not 0 <= level <= 8:
                raise ValueError(f"列表级别必须在 0-8 之间")

            doc = Document(str(file_path))

            # 添加列表段落
            if list_type == "bullet":
                paragraph = doc.add_paragraph(text, style='List Bullet')
            elif list_type == "number":
                paragraph = doc.add_paragraph(text, style='List Number')
            else:
                raise ValueError(f"不支持的列表类型: {list_type}")

            # 设置列表级别
            paragraph.paragraph_format.left_indent = Inches(0.5 * level)

            doc.save(str(file_path))

            logger.info(f"列表段落添加成功: {file_path}")
            return {
                "success": True,
                "message": "列表段落添加成功",
                "filename": str(file_path),
                "list_type": list_type,
                "level": level,
            }

        except Exception as e:
            logger.error(f"添加列表段落失败: {e}")
            return {"success": False, "message": f"添加失败: {str(e)}"}

    def add_multilevel_list(
        self,
        filename: str,
        items: list[dict[str, Any]],
        list_type: str = "bullet",
    ) -> dict[str, Any]:
        """添加多级列表.

        Args:
            filename: 文件名
            items: 列表项，每项包含 'text' 和 'level' (0-8)
                例如: [
                    {"text": "一级项目", "level": 0},
                    {"text": "二级项目", "level": 1},
                    {"text": "三级项目", "level": 2}
                ]
            list_type: 列表类型 ('bullet' 或 'number')

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            for item in items:
                text = item.get('text', '')
                level = item.get('level', 0)

                if not 0 <= level <= 8:
                    logger.warning(f"列表级别 {level} 超出范围，使用 0")
                    level = 0

                # 添加列表段落
                if list_type == "bullet":
                    paragraph = doc.add_paragraph(text, style='List Bullet')
                elif list_type == "number":
                    paragraph = doc.add_paragraph(text, style='List Number')
                else:
                    raise ValueError(f"不支持的列表类型: {list_type}")

                # 设置列表级别
                paragraph.paragraph_format.left_indent = Inches(0.5 * level)

            doc.save(str(file_path))

            logger.info(f"多级列表添加成功: {file_path}, 共 {len(items)} 项")
            return {
                "success": True,
                "message": f"成功添加 {len(items)} 个列表项",
                "filename": str(file_path),
                "list_type": list_type,
                "item_count": len(items)
            }

        except Exception as e:
            logger.error(f"添加多级列表失败: {e}")
            return {"success": False, "message": f"添加失败: {str(e)}"}

    def split_table_cell(
        self,
        filename: str,
        table_index: int,
        row: int,
        col: int,
        num_rows: int = 2,
        num_cols: int = 1,
    ) -> dict[str, Any]:
        """拆分表格单元格.

        注意: python-docx 对单元格拆分的支持有限，
        此方法通过插入新行列来模拟拆分效果。

        Args:
            filename: 文件名
            table_index: 表格索引
            row: 行索引
            col: 列索引
            num_rows: 拆分成的行数
            num_cols: 拆分成的列数

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            if row >= len(table.rows) or col >= len(table.columns):
                raise ValueError(f"单元格位置超出表格边界")

            # 获取原单元格的文本
            original_text = table.cell(row, col).text

            # python-docx 不直接支持单元格拆分
            # 这里提供一个替代方案的说明
            logger.warning(
                "python-docx 库不支持直接拆分单元格。"
                "建议使用合并单元格的逆操作，或重新设计表格结构。"
            )

            return {
                "success": False,
                "message": "python-docx 库不支持单元格拆分功能。"
                          "建议重新设计表格结构或使用 Microsoft Word 手动拆分。",
                "filename": str(file_path),
                "limitation": "python-docx library limitation"
            }

        except Exception as e:
            logger.error(f"拆分单元格失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def sort_table(
        self,
        filename: str,
        table_index: int,
        column_index: int,
        reverse: bool = False,
        has_header: bool = True,
    ) -> dict[str, Any]:
        """对表格按指定列排序.

        Args:
            filename: 文件名
            table_index: 表格索引
            column_index: 排序列索引
            reverse: 是否降序排序
            has_header: 是否有表头（第一行不参与排序）

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            if column_index >= len(table.columns):
                raise ValueError(f"列索引 {column_index} 超出范围")

            # 提取数据
            start_row = 1 if has_header else 0
            rows_data = []

            for i in range(start_row, len(table.rows)):
                row = table.rows[i]
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                rows_data.append((i, row_data))

            # 排序
            try:
                # 尝试数值排序
                rows_data.sort(
                    key=lambda x: float(x[1][column_index]) if x[1][column_index] else 0,
                    reverse=reverse
                )
            except ValueError:
                # 如果不是数值，使用字符串排序
                rows_data.sort(
                    key=lambda x: x[1][column_index],
                    reverse=reverse
                )

            # 重新填充数据
            for new_idx, (old_idx, row_data) in enumerate(rows_data):
                actual_row = new_idx + start_row
                for col_idx, cell_text in enumerate(row_data):
                    table.rows[actual_row].cells[col_idx].text = cell_text

            doc.save(str(file_path))

            logger.info(f"表格排序成功: {file_path}")
            return {
                "success": True,
                "message": f"成功按第 {column_index} 列排序",
                "filename": str(file_path),
                "column_index": column_index,
                "reverse": reverse,
                "sorted_rows": len(rows_data)
            }

        except Exception as e:
            logger.error(f"表格排序失败: {e}")
            return {"success": False, "message": f"排序失败: {str(e)}"}

    def import_table_data(
        self,
        filename: str,
        data: list[list[str]],
        has_header: bool = True,
        table_style: str = "Table Grid",
        insert_position: Optional[int] = None,
    ) -> dict[str, Any]:
        """从数据导入创建表格.

        Args:
            filename: 文件名
            data: 表格数据（二维列表）
            has_header: 第一行是否为表头
            table_style: 表格样式
            insert_position: 插入位置（段落索引，None表示在文档末尾）

        Returns:
            dict: 操作结果
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            if not data or not data[0]:
                raise ValueError("数据不能为空")

            doc = Document(str(file_path))

            rows = len(data)
            cols = len(data[0])

            # 创建表格
            if insert_position is None:
                # 在文档末尾添加表格
                table = doc.add_table(rows=rows, cols=cols)
            else:
                # 在指定位置插入表格
                if insert_position < 0 or insert_position >= len(doc.paragraphs):
                    raise ValueError(f"插入位置 {insert_position} 超出范围 (0-{len(doc.paragraphs)-1})")

                # 在指定段落之前插入一个新段落，然后在该段落后添加表格
                insert_para = doc.paragraphs[insert_position]
                new_para = insert_para.insert_paragraph_before()

                # 由于python-docx的限制，我们需要先在末尾创建表格，然后移动它
                # 这里使用一个技巧：在指定位置插入表格的XML元素
                table = doc.add_table(rows=rows, cols=cols)

                # 获取表格的XML元素
                tbl_element = table._element
                # 从当前位置移除
                tbl_element.getparent().remove(tbl_element)
                # 插入到指定位置
                new_para._element.addnext(tbl_element)

            table.style = table_style

            # 填充数据
            for i, row_data in enumerate(data):
                for j, cell_data in enumerate(row_data):
                    if j < cols:
                        cell = table.rows[i].cells[j]
                        cell.text = str(cell_data)

                        # 如果是表头，设置加粗
                        if has_header and i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

            doc.save(str(file_path))

            logger.info(f"表格数据导入成功: {file_path}, {rows}x{cols}, 插入位置: {insert_position}")
            return {
                "success": True,
                "message": f"成功导入 {rows}x{cols} 表格",
                "filename": str(file_path),
                "rows": rows,
                "cols": cols,
                "insert_position": insert_position,
            }

        except Exception as e:
            logger.error(f"导入表格数据失败: {e}")
            return {"success": False, "message": f"导入失败: {str(e)}"}
