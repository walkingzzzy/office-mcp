"""Word 表格格式化模块."""

from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger

from office_mcp_server.config import config
from office_mcp_server.utils.file_manager import FileManager
from office_mcp_server.utils.format_helper import ColorUtils


class WordTableFormatOperations:
    """Word 表格格式化操作类."""

    def __init__(self) -> None:
        """初始化表格格式化操作类."""
        self.file_manager = FileManager()

    def format_table_cell(
        self,
        filename: str,
        table_index: int,
        row: int,
        col: int,
        alignment: Optional[str] = None,
        background_color: Optional[str] = None,
        text_color: Optional[str] = None,
        bold: bool = False,
        font_size: Optional[int] = None,
    ) -> dict[str, Any]:
        """格式化表格单元格.

        Args:
            filename: 文件名
            table_index: 表格索引
            row: 行索引
            col: 列索引
            alignment: 对齐方式 ('left', 'center', 'right')
            background_color: 背景颜色 (HEX格式)
            text_color: 文字颜色 (HEX格式)
            bold: 是否加粗
            font_size: 字号
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            if row >= len(table.rows) or col >= len(table.columns):
                raise ValueError(f"单元格位置超出范围")

            cell = table.cell(row, col)

            # 设置对齐方式
            if alignment:
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                }
                if alignment in alignment_map:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = alignment_map[alignment]

            # 设置背景颜色
            if background_color:
                r, g, b = ColorUtils.hex_to_rgb(background_color)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), background_color.lstrip('#'))
                cell._element.get_or_add_tcPr().append(shading_elm)

            # 设置文字格式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if text_color:
                        r, g, b = ColorUtils.hex_to_rgb(text_color)
                        run.font.color.rgb = RGBColor(r, g, b)
                    if bold:
                        run.font.bold = True
                    if font_size:
                        run.font.size = Pt(font_size)

            doc.save(str(file_path))

            logger.info(f"单元格格式化成功: {file_path}")
            return {
                "success": True,
                "message": "单元格格式化成功",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"格式化单元格失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def apply_table_style(
        self,
        filename: str,
        table_index: int,
        style_name: str = "Table Grid",
    ) -> dict[str, Any]:
        """应用表格样式.

        Args:
            filename: 文件名
            table_index: 表格索引
            style_name: 样式名称
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            # 可用样式
            available_styles = [
                "Table Grid",
                "Light Shading",
                "Light List",
                "Light Grid",
                "Medium Shading 1",
                "Medium Shading 2",
                "Medium List 1",
                "Medium List 2",
                "Medium Grid 1",
                "Medium Grid 2",
                "Medium Grid 3",
                "Dark List",
                "Colorful Shading",
                "Colorful List",
                "Colorful Grid",
            ]

            if style_name not in available_styles:
                logger.warning(f"样式 '{style_name}' 可能不可用,将尝试应用")

            table.style = style_name

            doc.save(str(file_path))

            logger.info(f"表格样式应用成功: {file_path}")
            return {
                "success": True,
                "message": f"成功应用样式 '{style_name}'",
                "filename": str(file_path),
                "style_name": style_name,
                "available_styles": available_styles,
            }

        except Exception as e:
            logger.error(f"应用表格样式失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def set_table_borders(
        self,
        filename: str,
        table_index: int,
        border_style: str = "single",
        border_size: int = 4,
        border_color: str = "#000000",
    ) -> dict[str, Any]:
        """设置表格边框.

        Args:
            filename: 文件名
            table_index: 表格索引
            border_style: 边框样式 ('single', 'double', 'dotted', 'dashed')
            border_size: 边框粗细 (1-96)
            border_color: 边框颜色 (HEX格式)
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            # 设置表格边框
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            tblBorders = OxmlElement('w:tblBorders')

            # 边框位置
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), border_style)
                border.set(qn('w:sz'), str(border_size))
                border.set(qn('w:color'), border_color.lstrip('#'))
                tblBorders.append(border)

            tblPr.append(tblBorders)

            doc.save(str(file_path))

            logger.info(f"表格边框设置成功: {file_path}")
            return {
                "success": True,
                "message": "表格边框设置成功",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"设置表格边框失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def set_column_width(
        self,
        filename: str,
        table_index: int,
        col_index: int,
        width_inches: float,
    ) -> dict[str, Any]:
        """设置列宽.

        Args:
            filename: 文件名
            table_index: 表格索引
            col_index: 列索引
            width_inches: 列宽(英寸)
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            if col_index >= len(table.columns):
                raise ValueError(f"列索引 {col_index} 超出范围")

            # 设置列宽
            for row in table.rows:
                row.cells[col_index].width = Inches(width_inches)

            doc.save(str(file_path))

            logger.info(f"列宽设置成功: {file_path}")
            return {
                "success": True,
                "message": f"成功设置第 {col_index} 列宽度为 {width_inches} 英寸",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"设置列宽失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def set_row_height(
        self,
        filename: str,
        table_index: int,
        row_index: int,
        height_inches: float,
    ) -> dict[str, Any]:
        """设置行高.

        Args:
            filename: 文件名
            table_index: 表格索引
            row_index: 行索引
            height_inches: 行高(英寸)
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            if row_index >= len(table.rows):
                raise ValueError(f"行索引 {row_index} 超出范围")

            # 设置行高
            table.rows[row_index].height = Inches(height_inches)

            doc.save(str(file_path))

            logger.info(f"行高设置成功: {file_path}")
            return {
                "success": True,
                "message": f"成功设置第 {row_index} 行高度为 {height_inches} 英寸",
                "filename": str(file_path),
            }

        except Exception as e:
            logger.error(f"设置行高失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}

    def read_table_data(
        self,
        filename: str,
        table_index: int,
    ) -> dict[str, Any]:
        """读取表格数据.

        Args:
            filename: 文件名
            table_index: 表格索引
        """
        try:
            file_path = config.paths.output_dir / filename
            self.file_manager.validate_file_path(file_path, must_exist=True)

            doc = Document(str(file_path))

            if table_index >= len(doc.tables):
                raise ValueError(f"表格索引 {table_index} 超出范围")

            table = doc.tables[table_index]

            # 读取表格数据
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                data.append(row_data)

            logger.info(f"表格数据读取成功: {file_path}")
            return {
                "success": True,
                "message": "表格数据读取成功",
                "filename": str(file_path),
                "rows": len(table.rows),
                "columns": len(table.columns),
                "data": data,
            }

        except Exception as e:
            logger.error(f"读取表格数据失败: {e}")
            return {"success": False, "message": f"操作失败: {str(e)}"}
